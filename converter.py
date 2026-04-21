"""
格式转换模块：PDF / DOCX / SRT → TXT
拖入文件时自动调用，转换结果保存为同目录下的 .txt 文件
转换后自动经过 text_normalizer 美化管道。
"""
import re
from pathlib import Path
from text_normalizer import normalize as _normalize


# ── PDF → TXT ─────────────────────────────────────────────────────────────────

def _is_scanned_pdf(doc) -> bool:
    """检测是否为扫描版 PDF（无文字层）"""
    sample = min(5, doc.page_count)
    total_chars = sum(len(doc[i].get_text().strip()) for i in range(sample))
    return (total_chars / sample) < 50


def pdf_to_txt(pdf_path: str) -> str:
    """
    提取 PDF 文字层。若为扫描版则抛出 ValueError。
    返回纯文本字符串。
    """
    try:
        import fitz
    except ImportError:
        raise ImportError("需要安装 PyMuPDF：pip install pymupdf")

    doc = fitz.open(pdf_path)
    if _is_scanned_pdf(doc):
        doc.close()
        raise ValueError("此 PDF 为扫描版（无文字层），无法转换。\n请使用带文字层的 PDF。")

    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


# ── DOCX → TXT ────────────────────────────────────────────────────────────────

# 识别飞书妙记 / 讯飞听见 说话人行的正则
_SPEAKER_PATTERNS = [
    re.compile(r'^.{1,15}\s{2,}\d{2}:\d{2}:\d{2}'),   # 飞书：名字  00:00:05
    re.compile(r'^\[.{1,10}\]\d{2}:\d{2}:\d{2}'),       # 讯飞：[说话人]00:00:25
    re.compile(r'^\d{2}:\d{2}:\d{2}'),                  # 纯时间戳
]

def _is_speaker_line(text: str) -> bool:
    return any(p.match(text.strip()) for p in _SPEAKER_PATTERNS)


def docx_to_txt(docx_path: str) -> str:
    """
    提取 DOCX 纯文本，兼容飞书妙记 / 讯飞听见格式（说话人行被跳过）。
    返回纯文本字符串。
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx：pip install python-docx")

    doc = Document(docx_path)
    lines = []

    # 普通段落
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or _is_speaker_line(t):
            continue
        lines.append(t)

    # 表格（飞书部分格式用表格存说话人|内容）
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and _is_speaker_line(cells[0]):
                # 说话人列跳过，只取内容列
                content = cells[1].strip()
                if content:
                    lines.append(content)
            else:
                for c in cells:
                    if c and not _is_speaker_line(c):
                        lines.append(c)

    return "\n".join(lines)


# ── SRT → TXT ─────────────────────────────────────────────────────────────────

_SRT_PATTERN = re.compile(
    r'^\d+\s*\n'
    r'\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}\s*\n'
    r'([\s\S]*?)(?=\n\d+\s*\n|\Z)',
    re.MULTILINE
)

_SRT_BLOCK_RE = re.compile(
    r'^\d+\s*\n'
    r'(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\s*\n'
    r'([\s\S]*?)(?=\n\d+\s*\n|\Z)',
    re.MULTILINE
)

# 匹配所有非汉字、非字母数字、非空白的字符（即所有标点）
_READING_PUNCT_RE = re.compile(
    r'[^一-鿿㐀-䶿豈-﫿A-Za-z0-9\s]'
)

def _srt_ms(t: str) -> int:
    """SRT 时间戳 HH:MM:SS,mmm → 毫秒"""
    h, m, rest = t.split(':')
    s, ms = rest.split(',')
    return int(h) * 3_600_000 + int(m) * 60_000 + int(s) * 1_000 + int(ms)


def srt_to_txt(srt_path: str, reading_mode: bool = True) -> str:
    """
    解析 SRT 字幕文件，去掉序号和时间戳，返回纯文本。
    reading_mode=True（默认）：按时间间隔分段、去除所有标点，适合阅读排版。
    reading_mode=False：保留标点，每条字幕一行。
    自动处理 UTF-8 / UTF-8-BOM / GBK 编码。
    """
    content = None
    for enc in ('utf-8-sig', 'utf-8', 'gbk', 'latin-1'):
        try:
            content = Path(srt_path).read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    if content is None:
        raise ValueError("无法识别 SRT 文件编码")

    content = content.replace('\r\n', '\n').strip()

    if not reading_mode:
        lines = []
        for match in _SRT_PATTERN.finditer(content + '\n\n'):
            text = match.group(1).strip().replace('\n', ' ')
            if text:
                lines.append(text)
        return "\n".join(lines)

    # 阅读排版模式：解析时间戳，按停顿间隔（>1.5s）分段
    blocks = []
    for m in _SRT_BLOCK_RE.finditer(content + '\n\n'):
        start_ms = _srt_ms(m.group(1))
        end_ms   = _srt_ms(m.group(2))
        text = m.group(3).strip().replace('\n', ' ')
        if text:
            blocks.append((start_ms, end_ms, text))

    if not blocks:
        return ''

    GAP_THRESHOLD_MS = 1500
    groups: list[list[str]] = []
    current: list[str] = [blocks[0][2]]
    prev_end = blocks[0][1]

    for start_ms, end_ms, text in blocks[1:]:
        if start_ms - prev_end > GAP_THRESHOLD_MS:
            groups.append(current)
            current = []
        current.append(text)
        prev_end = end_ms
    groups.append(current)

    result: list[str] = []
    for group in groups:
        para = ' '.join(group)
        para = _READING_PUNCT_RE.sub(' ', para)
        para = re.sub(r'[ \t]{2,}', ' ', para).strip()
        if para:
            result.append(para)

    return '\n\n'.join(result)


# ── 统一入口 ──────────────────────────────────────────────────────────────────

SUPPORTED_EXTS = {'.pdf', '.docx', '.srt'}

def convert_to_txt(src_path: str) -> str:
    """
    将 PDF / DOCX / SRT 转换并保存为同目录下的 .txt 文件。
    返回生成的 .txt 文件路径。
    若已存在同名 .txt 则直接覆盖。
    """
    src = Path(src_path)
    ext = src.suffix.lower()

    if ext == '.pdf':
        text = pdf_to_txt(src_path)
    elif ext == '.docx':
        text = docx_to_txt(src_path)
    elif ext == '.srt':
        text = srt_to_txt(src_path, reading_mode=True)
    else:
        raise ValueError(f"不支持的格式：{ext}，仅支持 .pdf / .docx / .srt")

    # SRT 阅读模式已完成排版，不能再跑 Stage 5（会删除段落间空格）
    if ext == '.srt':
        text = text.strip()
    else:
        text = _normalize(
            text,
            reconstruct_paragraphs=True,
            normalize_punct=True,
            pangu_spacing=False,
            clean_fillers=False,
        )

    dst = src.with_suffix('.txt')
    dst.write_text(text, encoding='utf-8')
    return str(dst)
