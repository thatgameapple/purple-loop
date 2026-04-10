#!/usr/bin/env python3
"""逐字稿阅读器"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, font as tkfont
import os, json, uuid, re, subprocess, shutil, platform
from annotation_manager import AnnotationManager
import theme as _theme
try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    _HAS_DND = True
except ImportError:
    _HAS_DND = False

DATA_FILE = os.path.expanduser('~/.script_reader_data.json')
FONTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fonts')

# ── 平台检测 ──────────────────────────────────────────────────────
IS_MAC  = platform.system() == 'Darwin'
IS_WIN  = platform.system() == 'Windows'
MOD     = 'Command' if IS_MAC else 'Control'   # tkinter 绑定用
MOD_KEY = 'Cmd'     if IS_MAC else 'Ctrl'      # 菜单显示用
UI_FONT = 'PingFang SC' if IS_MAC else 'Microsoft YaHei'  # UI 字体

# 内置字体：文件名 → (显示名, tkinter 字族名)
BUILTIN_FONTS = {
    'LXGWWenKai-Light.ttf':   ('霞鹜文楷 Light',   'LXGW WenKai'),
    'LXGWWenKai-Regular.ttf': ('霞鹜文楷 Regular', 'LXGW WenKai'),
    'LXGWWenKai-Medium.ttf':  ('霞鹜文楷 Medium',  'LXGW WenKai'),
}

def _install_builtin_fonts():
    """将内置字体安装到系统字体目录（首次运行时执行）"""
    if IS_MAC:
        dest_dir = os.path.expanduser('~/Library/Fonts')
    elif IS_WIN:
        dest_dir = os.path.join(os.environ.get('LOCALAPPDATA', ''),
                                'Microsoft', 'Windows', 'Fonts')
    else:
        dest_dir = os.path.expanduser('~/.local/share/fonts')
    os.makedirs(dest_dir, exist_ok=True)
    for fname in BUILTIN_FONTS:
        src = os.path.join(FONTS_DIR, fname)
        dst = os.path.join(dest_dir, fname)
        if os.path.exists(src) and not os.path.exists(dst):
            try:
                shutil.copy2(src, dst)
            except Exception:
                pass

# ── 色板（从 theme 模块取可变引用，主题切换时自动同步） ──────────
C = _theme.C


class TagStore:
    def __init__(self, path):
        self.path = path
        self.tags = {}
        self.files = {}
        self._load()

    def _load(self):
        if os.path.exists(self.path):
            try:
                data = json.loads(open(self.path, encoding='utf-8').read())
                self.tags  = data.get('tags',  {})
                self.files = data.get('files', {})
            except Exception:
                pass

    def save(self):
        open(self.path, 'w', encoding='utf-8').write(
            json.dumps({'tags': self.tags, 'files': self.files},
                       ensure_ascii=False, indent=2))

    def add_tag(self, name, parent_id=None):
        tid = str(uuid.uuid4())
        self.tags[tid] = {'name': name, 'parent': parent_id, 'children': []}
        if parent_id and parent_id in self.tags:
            self.tags[parent_id]['children'].append(tid)
        self.save()
        return tid

    def rename_tag(self, tid, name):
        if tid in self.tags:
            self.tags[tid]['name'] = name
            self.save()

    def delete_tag(self, tid):
        if tid not in self.tags:
            return
        for c in list(self.tags[tid]['children']):
            self.delete_tag(c)
        pid = self.tags[tid]['parent']
        if pid and pid in self.tags:
            self.tags[pid]['children'] = [x for x in self.tags[pid]['children'] if x != tid]
        for fp in list(self.files.keys()):
            if isinstance(self.files[fp], list):
                self.files[fp] = [x for x in self.files[fp] if x != tid]
        del self.tags[tid]
        self.save()

    def get_roots(self):
        return [tid for tid, t in self.tags.items() if not t['parent']]

    def get_pinned(self):
        return [tid for tid in self.get_roots() if self.tags[tid].get('pinned')]

    def pin_tag(self, tid, pinned=True):
        if tid in self.tags:
            self.tags[tid]['pinned'] = pinned
            self.save()

    def add_file(self, fp, tag_id=None):
        if fp not in self.files:
            self.files[fp] = []
        if tag_id and tag_id not in self.files[fp]:
            self.files[fp].append(tag_id)
        self.save()

    def remove_file(self, fp, tag_id):
        if fp in self.files:
            self.files[fp] = [x for x in self.files[fp] if x != tag_id]
            if not self.files[fp]:
                del self.files[fp]
            self.save()

    def files_for(self, tag_id):
        return [fp for fp, tags in self.files.items()
                if not fp.startswith('__') and isinstance(tags, list) and tag_id in tags]

    def count_under(self, tid):
        t = self.tags.get(tid, {})
        return len(self.files_for(tid)) + sum(self.count_under(c) for c in t.get('children', []))

    def _is_special(self, key):
        return key.startswith('__')


class App(TkinterDnD.Tk if _HAS_DND else tk.Tk):
    def __init__(self):
        super().__init__()
        self.withdraw()   # 启动期间隐藏，避免用户看到欢迎页闪烁
        self.title("逐字稿")
        self.geometry("1200x760")
        self.minsize(1000, 600)
        self.configure(bg=C['bg'])

        # 应用图标（可切换，store 初始化后调用）
        self._app_dir = os.path.dirname(os.path.abspath(__file__))
        self._icon_ref = None  # 防止被 GC

        _install_builtin_fonts()

        self.store = TagStore(DATA_FILE)
        self._load_icon()   # store 就绪后加载图标
        prefs = self.store.files.get('__prefs__', {})
        if not isinstance(prefs, dict):
            prefs = {}
        self._font_size = prefs.get('font_size', 18)
        self._font_family = prefs.get('font_family', 'LXGW WenKai')
        self._line_spacing = prefs.get('line_spacing', 10)  # iA Writer ~1.5x 行高
        self._pending_cb = None
        self._current_file = None
        self._chars_text   = ''
        self._autosave_after = None   # 自动保存定时器
        self._wordcount_after = None  # 字数统计防抖定时器
        self._resetting_modified = False  # 防止 <<Modified>> 重入
        self._search_matches = []
        self._search_idx = -1
        self._search_timer = None
        self._sidebar_visible = True
        self.annot_mgr = None   # 延迟到 UI 构建完成后初始化

        # 阅读计时
        self._reading_start = None   # 本次聚焦开始时间

        self._build_menu()
        self._build_ui()
        self._refresh_tree()
        self._refresh_author_lbl()
        # 标注系统（text widget 已就绪后初始化）
        self.annot_mgr = AnnotationManager(self, self.text, self.store, self._note_bar)
        self.annot_mgr.setup()
        self._configure_heading_tags()
        self.protocol('WM_DELETE_WINDOW', self._on_close)
        self.bind('<FocusIn>',   self._on_focus_in)
        self.bind('<FocusOut>',  self._on_focus_out)
        self.bind('<Configure>', self._on_window_resize)
        # 拖放文件
        if _HAS_DND:
            self.drop_target_register(DND_FILES)
            self.dnd_bind('<<Drop>>', self._on_file_drop)

        # 恢复上次保存的主题（在菜单构建完成后调用）
        if self.store.files.get('__follow_sys__', False):
            self._follow_sys.set(True)
            # 启动时始终强制应用系统主题，不做"是否相同"比较
            sys_theme = self._get_system_theme()
            self._apply_theme(sys_theme)
            self._sys_poll_id = self.after(2000, self._sync_system_theme)
        else:
            saved_theme = self.store.files.get('__theme__', 'dark')
            if saved_theme != 'dark':
                self._apply_theme(saved_theme)
            else:
                self._theme_var.set('dark')
        # 恢复图标菜单勾选
        self._icon_var.set(self.store.files.get('__icon__', 'zz'))

        # 恢复上次打开的文件（有文件记录则直接进阅读页，否则保留欢迎页）
        last = self.store.files.get('__last_file__', '')
        if last and os.path.isfile(last):
            self._load(last)

        self.deiconify()  # 一切就绪后再显示窗口，无闪烁
        self.lift()
        self.focus_force()

    def _build_menu(self):
        mb = tk.Menu(self)
        self.configure(menu=mb)

        fm = tk.Menu(mb, tearoff=0)
        mb.add_cascade(label="文件", menu=fm)
        fm.add_command(label="新建文稿…",    accelerator=f"{MOD_KEY}+N", command=self._new_doc)
        fm.add_separator()
        fm.add_command(label="导入 TXT…",    accelerator=f"{MOD_KEY}+O", command=self._import)
        fm.add_command(label="导入 Word…",   command=self._import_word)
        fm.add_command(label="导入 PDF…",    command=self._import_pdf)
        fm.add_separator()
        fm.add_command(label="保存",         accelerator=f"{MOD_KEY}+S", command=self._save)
        ex = tk.Menu(fm, tearoff=0)
        fm.add_cascade(label="导出", menu=ex)
        ex.add_command(label="导出为 PDF…",         command=self._export_pdf)
        ex.add_command(label="导出为 Word…",        command=self._export_word)
        ex.add_separator()
        ex.add_command(label="导出批注为 Markdown…", command=self._export_annots_md)
        ex.add_separator()
        ex.add_command(label="导出工程文件…",        command=self._export_zzg)
        fm.add_separator()
        fm.add_command(label="导入工程文件…",      command=self._import_zzg)
        fm.add_separator()
        fm.add_command(label="撤销", accelerator=f"{MOD_KEY}+Z",
                       command=lambda: self.text.edit_undo())
        fm.add_command(label="重做", accelerator=f"{MOD_KEY}+Shift+Z",
                       command=lambda: self.text.edit_redo())
        fm.add_separator()
        fm.add_command(label="退出", command=self._on_close)

        vm = tk.Menu(mb, tearoff=0)
        mb.add_cascade(label="显示", menu=vm)
        vm.add_command(label="放大字号", accelerator=f"{MOD_KEY}+=", command=lambda: self._font_change(1))
        vm.add_command(label="缩小字号", accelerator=f"{MOD_KEY}+−", command=lambda: self._font_change(-1))
        vm.add_separator()
        vm.add_command(label="增大行距", accelerator=f"{MOD_KEY}+]", command=lambda: self._spacing_change(2))
        vm.add_command(label="减小行距", accelerator=f"{MOD_KEY}+[", command=lambda: self._spacing_change(-2))
        vm.add_separator()
        vm.add_command(label="选择字体…", command=self._show_font_picker)
        vm.add_command(label="导入字体文件…", command=self._import_font_file)
        vm.add_separator()
        # 主题切换
        self._theme_var   = tk.StringVar(value='dark')
        self._follow_sys  = tk.BooleanVar(value=False)
        self._sys_poll_id = None   # after() 轮询 ID
        vm.add_radiobutton(label="护眼模式", variable=self._theme_var,
                           value='dark',  command=lambda: self._set_theme_manual('dark'))
        vm.add_radiobutton(label="白天模式", variable=self._theme_var,
                           value='light', command=lambda: self._set_theme_manual('light'))
        vm.add_checkbutton(label="跟随系统", variable=self._follow_sys,
                           command=self._toggle_follow_system)
        vm.add_separator()
        # 图标选择
        self._icon_var = tk.StringVar(value='zz')
        vm.add_radiobutton(label="图标：Zz", variable=self._icon_var,
                           value='zz', command=lambda: self._switch_icon('zz'))
        vm.add_radiobutton(label="图标：雨中漫步", variable=self._icon_var,
                           value='custom', command=lambda: self._switch_icon('custom'))
        vm.add_separator()
        vm.add_command(label="显示/隐藏侧栏", accelerator=f"{MOD_KEY}+\\", command=self._toggle_sidebar)
        vm.add_command(label="搜索", accelerator=f"{MOD_KEY}+F", command=self._toggle_search)
        vm.add_command(label="全局搜索…", accelerator=f"{MOD_KEY}+Shift+F", command=self._open_global_search)
        vm.add_separator()
        vm.add_command(label="阅读统计…", command=self._show_stats)
        vm.add_separator()
        vm.add_command(label="设置作者名…", command=self._set_author)

        # ── 标注菜单 ──
        am = tk.Menu(mb, tearoff=0)
        mb.add_cascade(label="标注", menu=am)
        am.add_command(label="黄色高亮",  accelerator=f"{MOD_KEY}+1",
                       command=lambda: self.annot_mgr and self.annot_mgr._annotate_selection('hl_yellow'))
        am.add_command(label="绿色高亮",  accelerator=f"{MOD_KEY}+2",
                       command=lambda: self.annot_mgr and self.annot_mgr._annotate_selection('hl_green'))
        am.add_command(label="粉色高亮",  accelerator=f"{MOD_KEY}+3",
                       command=lambda: self.annot_mgr and self.annot_mgr._annotate_selection('hl_pink'))
        am.add_command(label="紫色高亮",  accelerator=f"{MOD_KEY}+4",
                       command=lambda: self.annot_mgr and self.annot_mgr._annotate_selection('hl_purple'))
        am.add_separator()
        am.add_command(label="加粗",      accelerator=f"{MOD_KEY}+B",
                       command=lambda: self.annot_mgr and self.annot_mgr._annotate_selection('bold'))
        am.add_command(label="下划线",    accelerator=f"{MOD_KEY}+U",
                       command=lambda: self.annot_mgr and self.annot_mgr._annotate_selection('underline'))
        am.add_separator()
        am.add_command(label="取消光标处标注", accelerator=f"{MOD_KEY}+Delete",
                       command=lambda: self.annot_mgr and self.annot_mgr._remove_at_cursor())
        am.add_separator()
        am.add_command(label="显示/隐藏标注面板", accelerator=f"{MOD_KEY}+Shift+A",
                       command=lambda: self.annot_mgr and self.annot_mgr.toggle_panel())

        tm = tk.Menu(mb, tearoff=0)
        mb.add_cascade(label="标签", menu=tm)
        tm.add_command(label="新建标签…", accelerator="Cmd+N", command=self._new_root_tag)

        hm = tk.Menu(mb, tearoff=0)
        mb.add_cascade(label="帮助", menu=hm)
        hm.add_command(label="使用说明…", command=self._show_help)

        self.bind_all('<Command-o>', lambda e: self._import())
        self.bind_all('<Command-n>', lambda e: self._new_doc())
        self.bind_all('<Command-s>', lambda e: self._save())
        self.bind_all('<Command-f>', lambda e: self._toggle_search())
        self.bind_all('<Command-F>', lambda e: self._open_global_search())
        self.bind_all('<Command-backslash>', lambda e: self._toggle_sidebar())
        self.bind_all('<Command-equal>', lambda e: self._font_change(1))
        self.bind_all('<Command-minus>', lambda e: self._font_change(-1))
        self.bind_all('<Command-bracketright>', lambda e: self._spacing_change(2))
        self.bind_all('<Command-bracketleft>', lambda e: self._spacing_change(-2))
        self.bind_all('<Command-A>',
                      lambda e: self.annot_mgr and self.annot_mgr.toggle_panel())

    def _build_ui(self):
        # ── 侧栏 ──────────────────────────────────────────
        self._sidebar = tk.Frame(self, bg=C['bg_sidebar'], width=230)
        self._sidebar.pack(side=tk.LEFT, fill=tk.Y)
        self._sidebar.pack_propagate(False)

        # ── 右侧标注面板（常驻，初始隐藏）────────────────
        self._annot_panel_frame = tk.Frame(self, bg=C['bg_sidebar'], width=240)
        self._annot_panel_frame.pack_propagate(False)
        # 初始不 pack，等 AnnotationManager 初始化后通过 toggle 显示
        self._annot_panel_visible = False

        # 侧栏顶部：Tab 切换（标签 / 目录）
        sb_top = tk.Frame(self._sidebar, bg=C['bg_sidebar'], height=48)
        sb_top.pack(fill=tk.X)
        sb_top.pack_propagate(False)

        self._sb_tab = tk.StringVar(value='tags')   # 'tags' | 'toc'

        def _make_tab(text, val):
            lbl = tk.Label(sb_top, text=text, bg=C['bg_sidebar'],
                           font=(UI_FONT, 11), cursor='hand2', padx=12, pady=4)
            def _upd():
                active = self._sb_tab.get() == val
                lbl.config(fg=C['accent'] if active else C['fg_dim'])
            lbl._upd = _upd
            lbl.bind('<Button-1>', lambda e, v=val: self._switch_sidebar_tab(v))
            lbl.pack(side=tk.LEFT, padx=(6, 0))
            _upd()
            return lbl

        self._tab_lbl_tags = _make_tab("标签", 'tags')
        self._tab_lbl_toc  = _make_tab("目录", 'toc')

        plus_lbl = tk.Label(sb_top, text="＋", bg=C['bg_sidebar'], fg=C['fg_dim'],
                             font=(UI_FONT, 16), cursor='hand2')
        plus_lbl.pack(side=tk.RIGHT, padx=14)
        plus_lbl.bind('<Button-1>', lambda e: self._new_root_tag())
        plus_lbl.bind('<Enter>', lambda e: plus_lbl.config(fg=C['accent']))
        plus_lbl.bind('<Leave>', lambda e: plus_lbl.config(fg=C['fg_dim']))


        # Treeview 样式
        s = ttk.Style()
        s.theme_use('default')
        s.configure('Sidebar.Treeview',
                    background=C['bg_sidebar'], foreground=C['fg_tag'],
                    fieldbackground=C['bg_sidebar'], rowheight=36, borderwidth=0,
                    indent=16)
        s.map('Sidebar.Treeview',
              background=[('selected', C['bg_sel_tag'])],
              foreground=[('selected', C['accent'])])   # Flomo：选中态用强调色
        s.configure('Sidebar.Treeview.Heading', background=C['bg_sidebar'])

        # 侧栏滚动条（极细）
        s.configure('Thin.Vertical.TScrollbar',
                    troughcolor=C['bg_sidebar'], background=C['accent_dim'],
                    width=4, relief=tk.FLAT, borderwidth=0)
        s.map('Thin.Vertical.TScrollbar',
              background=[('active', C['accent'])])

        # ── 侧栏底栏：作者身份 ───────────────────────────────
        sb_bottom = tk.Frame(self._sidebar, bg=C['bg_sidebar'])
        sb_bottom.pack(side=tk.BOTTOM, fill=tk.X)
        tk.Frame(sb_bottom, bg=C['border'], height=1).pack(fill=tk.X)

        author_row = tk.Frame(sb_bottom, bg=C['bg_sidebar'])
        author_row.pack(fill=tk.X, padx=16, pady=10)

        self._author_lbl = tk.Label(author_row, text="",
                                     bg=C['bg_sidebar'], fg=C['fg_file'],
                                     font=(UI_FONT, 11), cursor='hand2',
                                     anchor='w')
        self._author_lbl.pack(side=tk.LEFT, fill=tk.X, expand=True)

        def _refresh_author_lbl():
            name = self._get_author()
            if name:
                self._author_lbl.config(text=name, fg=C['accent'])
            else:
                self._author_lbl.config(text="点击设置作者名", fg=C['fg_hint'])

        self._refresh_author_lbl = _refresh_author_lbl

        def _on_author_click(e):
            self._set_author()

        self._author_lbl.bind('<Button-1>', _on_author_click)
        self._author_lbl.bind('<Enter>',
            lambda e: self._author_lbl.config(fg=C['accent']))
        self._author_lbl.bind('<Leave>',
            lambda e: self._author_lbl.config(
                fg=C['accent'] if self._get_author() else C['fg_hint']))

        # 用一个容器叠放 标签/目录 两个面板，通过 tkraise 切换，避免 pack 闪烁
        _tab_container = tk.Frame(self._sidebar, bg=C['bg_sidebar'])
        _tab_container.pack(fill=tk.BOTH, expand=True)
        self._tab_container = _tab_container

        _tree_frame = tk.Frame(_tab_container, bg=C['bg_sidebar'])
        _tree_frame.place(relx=0, rely=0, relwidth=1, relheight=1)
        self._tree_frame_ref = _tree_frame

        _vsb_tree = ttk.Scrollbar(_tree_frame, orient=tk.VERTICAL,
                                   style='Thin.Vertical.TScrollbar')
        _vsb_tree.pack(side=tk.RIGHT, fill=tk.Y)

        self.tree = ttk.Treeview(_tree_frame, show='tree',
                                  selectmode='browse', style='Sidebar.Treeview',
                                  yscrollcommand=_vsb_tree.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        _vsb_tree.config(command=self.tree.yview)

        # 支持鼠标滚轮滚动侧栏
        self.tree.bind('<MouseWheel>',
                       lambda e: self.tree.yview_scroll(int(-1*(e.delta/120)), 'units'))
        self.tree.tag_configure('tag',    foreground=C['fg_tag'])
        self.tree.tag_configure('file',   foreground=C['fg_file'])
        self.tree.tag_configure('header', foreground=C['fg_dim2'])
        self.tree.bind('<<TreeviewSelect>>', self._on_tree_select)
        self.tree.bind('<<TreeviewOpen>>',   self._on_tree_toggle)
        self.tree.bind('<<TreeviewClose>>', self._on_tree_toggle)
        self.tree.bind('<Button-2>', self._on_rclick)
        self.tree.bind('<Button-3>', self._on_rclick)

        # Tooltip：悬停显示完整文件名
        self._tooltip_data  = {}   # item_id -> 完整显示文字
        self._tooltip_win   = None
        self._tooltip_item  = None
        self.tree.bind('<Motion>', self._on_tree_motion)
        self.tree.bind('<Leave>',  self._hide_tooltip)

        self._ctx = tk.Menu(self, tearoff=0)

        # ── TOC 面板（与标签面板叠放，通过 tkraise 切换）──────────
        self._toc_frame = tk.Frame(self._tab_container, bg=C['bg_sidebar'])
        self._toc_frame.place(relx=0, rely=0, relwidth=1, relheight=1)

        _vsb_toc = ttk.Scrollbar(self._toc_frame, orient=tk.VERTICAL,
                                  style='Thin.Vertical.TScrollbar')
        _vsb_toc.pack(side=tk.RIGHT, fill=tk.Y)

        self._toc_tree = ttk.Treeview(self._toc_frame, show='tree',
                                       selectmode='browse', style='Sidebar.Treeview',
                                       yscrollcommand=_vsb_toc.set)
        self._toc_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        _vsb_toc.config(command=self._toc_tree.yview)
        self._toc_tree.bind('<MouseWheel>',
            lambda e: self._toc_tree.yview_scroll(int(-1*(e.delta/120)), 'units'))
        self._toc_tree.bind('<<TreeviewSelect>>', self._on_toc_select)
        self._toc_tree.bind('<Motion>', self._on_toc_motion)
        self._toc_tree.bind('<Leave>',  self._hide_tooltip)

        # 确保标签面板默认在最上面
        self._tree_frame_ref.tkraise()

        # ── 阅读区（注意：必须在右侧面板 pack 之后再 pack，保证面板占右侧）──
        self._reader = tk.Frame(self, bg=C['bg'])
        self._reader.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 顶栏（三区：左=侧栏切换 / 中=文件名居中 / 右=字数·进度）
        topbar = tk.Frame(self._reader, bg=C['bg'], height=44)
        topbar.pack(fill=tk.X)
        topbar.pack_propagate(False)

        # 左区：侧栏折叠按钮
        _sb_toggle = tk.Label(topbar, text="☰", bg=C['bg'], fg=C['fg_dim'],
                              font=(UI_FONT, 13), cursor='hand2', padx=16)
        _sb_toggle.pack(side=tk.LEFT)
        _sb_toggle.bind('<Button-1>', lambda e: self._toggle_sidebar())
        _sb_toggle.bind('<Enter>', lambda e: _sb_toggle.config(fg=C['fg']))
        _sb_toggle.bind('<Leave>', lambda e: _sb_toggle.config(fg=C['fg_dim']))

        # 右区：字数 · 进度（右侧先 pack，确保中间 label 能正确居中）
        self._lbl_meta = tk.Label(topbar, text="",
                                   bg=C['bg'], fg=C['fg_dim'],
                                   font=(UI_FONT, 11))
        self._lbl_meta.pack(side=tk.RIGHT, padx=20)

        self._lbl_time = tk.Label(topbar, text="",
                                   bg=C['bg'], fg=C['fg_hint'],
                                   font=(UI_FONT, 11))
        self._lbl_time.pack(side=tk.RIGHT, padx=(0, 4))

        # 中区：文件名居中
        self._lbl_file = tk.Label(topbar, text="",
                                   bg=C['bg'], fg=C['fg_dim'],
                                   font=(UI_FONT, 12))
        self._lbl_file.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 兼容旧代码（_lbl_chars 某些地方被引用）
        self._lbl_chars = self._lbl_meta

        # 极细分割线
        tk.Frame(self._reader, bg=C['border'], height=1).pack(fill=tk.X)


        # 搜索栏（默认隐藏）
        self._search_bar = tk.Frame(self._reader, bg=C['bg_input'], pady=10)
        s_inner = tk.Frame(self._search_bar, bg=C['bg_input'])
        s_inner.pack(fill=tk.X, padx=24)

        # 输入框容器（带高亮边框）
        input_wrap = tk.Frame(s_inner, bg=C['bg_input'],
                              highlightbackground=C['accent_dim'],
                              highlightthickness=1)
        input_wrap.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 放大镜图标
        tk.Label(input_wrap, text="⌕", bg=C['bg_input'], fg=C['fg_hint'],
                 font=('PingFang SC', 16), padx=10).pack(side=tk.LEFT)

        self._search_var = tk.StringVar()
        self._search_entry = tk.Entry(input_wrap, textvariable=self._search_var,
                                       bg=C['bg_input'], fg=C['fg'],
                                       insertbackground=C['fg'],
                                       font=('PingFang SC', 13),
                                       relief=tk.FLAT, bd=0)
        self._search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=7)
        self._search_entry.bind('<Return>',       lambda e: self._search_next())
        self._search_entry.bind('<Shift-Return>', lambda e: self._search_prev())
        self._search_entry.bind('<Escape>',       lambda e: self._close_search())
        self._search_var.trace_add('write', self._search_changed)

        # 聚焦时高亮边框
        def _focus_in(_e):
            input_wrap.config(highlightbackground=C['accent'])
        def _focus_out(_e):
            input_wrap.config(highlightbackground=C['accent_dim'])
        self._search_entry.bind('<FocusIn>',  _focus_in)
        self._search_entry.bind('<FocusOut>', _focus_out)

        # 右侧控件区
        ctrl = tk.Frame(s_inner, bg=C['bg_input'])
        ctrl.pack(side=tk.LEFT, padx=(10, 0))

        # 计数标签
        self._search_count_lbl = tk.Label(ctrl, text="",
                                           bg=C['bg_input'], fg=C['accent'],
                                           font=('PingFang SC', 11), width=7,
                                           anchor='e')
        self._search_count_lbl.pack(side=tk.LEFT)

        # ↑ ↓ 导航按钮
        for txt, cmd in [("↑", self._search_prev), ("↓", self._search_next)]:
            nav = tk.Label(ctrl, text=txt, bg=C['bg_input'], fg=C['fg_file'],
                           font=('PingFang SC', 13), padx=9, pady=4,
                           cursor='hand2',
                           highlightbackground=C['accent_dim'],
                           highlightthickness=1)
            nav.pack(side=tk.LEFT, padx=(6, 0))
            nav.bind('<Button-1>', lambda e, c=cmd: c())
            nav.bind('<Enter>',
                     lambda e, b=nav: b.config(fg=C['fg'], bg=C['accent_dim'],
                                               highlightbackground=C['accent']))
            nav.bind('<Leave>',
                     lambda e, b=nav: b.config(fg=C['fg_file'], bg=C['bg_input'],
                                               highlightbackground=C['accent_dim']))

        # 关闭按钮
        close_s = tk.Label(ctrl, text="✕", bg=C['bg_input'], fg=C['btn_close'],
                           font=('PingFang SC', 13), padx=10, cursor='hand2')
        close_s.pack(side=tk.LEFT, padx=(8, 0))
        close_s.bind('<Button-1>', lambda e: self._close_search())
        close_s.bind('<Enter>', lambda e: close_s.config(fg=C['fg_file']))
        close_s.bind('<Leave>', lambda e: close_s.config(fg=C['btn_close']))

        # 极细分割线（搜索栏底部，仅搜索时可见）
        self._search_divider = tk.Frame(self._reader, bg=C['border'], height=1)

        # ── 文本区 ────────────────────────────────────────
        self.text = tk.Text(self._reader,
                            bg=C['bg'], fg=C['fg'],
                            font=(self._font_family, self._font_size),
                            wrap=tk.WORD, relief=tk.FLAT, bd=0,
                            highlightthickness=0,
                            padx=72, pady=52,
                            spacing1=2, spacing2=self._line_spacing, spacing3=self._line_spacing * 2,
                            cursor='arrow',
                            insertbackground=C['fg'],
                            selectbackground=C['select_bg'],
                            selectforeground=C['fg'],
                            undo=True, maxundo=-1)

        # 搜索高亮
        self.text.tag_configure('search_match',
                                background=C['hl_all_bg'], foreground=C['hl_all_fg'])
        self.text.tag_configure('search_current',
                                background=C['hl_cur_bg'], foreground=C['hl_cur_fg'])

        # 右键菜单
        self._text_ctx = tk.Menu(self, tearoff=0)
        self._text_ctx.add_command(label="复制", command=lambda: self.text.event_generate('<<Copy>>'))
        self._text_ctx.add_command(label="粘贴", command=lambda: self.text.event_generate('<<Paste>>'))
        self._text_ctx.add_separator()
        self._text_ctx.add_command(label="撤销", command=lambda: self.text.edit_undo())
        self._text_ctx.add_command(label="重做", command=lambda: self.text.edit_redo())
        # 标注选项将在弹出时动态填充（见 _text_rclick）
        self.text.bind('<Button-2>', self._text_rclick)
        self.text.bind('<Button-3>', self._text_rclick)
        self.text.bind('<KeyRelease>', self._on_key_release)
        self.text.bind('<Return>', self._check_notion_heading)


        # 文本区滚动条（极细）
        s.configure('Thin.Vertical.TScrollbar',
                    troughcolor=C['bg'], background=C['accent_dim'],
                    width=4, relief=tk.FLAT, borderwidth=0)

        # 底部备注栏（默认隐藏）
        self._note_bar = tk.Frame(self._reader, bg=C['note_bar'], height=130)

        # 底部进度条（5px 细条，无外框，必须在文本区之前 pack）
        self._progress_track = tk.Frame(self._reader, bg=C['border'], height=5)
        self._progress_track.pack(side=tk.BOTTOM, fill=tk.X)
        self._progress_fill = tk.Frame(self._progress_track, bg=C['accent'], height=5)
        self._progress_fill.place(x=0, y=0, relheight=1, width=0)

        self._vsb_text = ttk.Scrollbar(self._reader, orient=tk.VERTICAL,
                                        command=self.text.yview,
                                        style='Thin.Vertical.TScrollbar')
        self.text.configure(yscrollcommand=self._on_text_scroll)
        self._vsb_text.pack(side=tk.RIGHT, fill=tk.Y)

        # ── 空状态（无文件时显示，加载文件后隐藏）────────────
        self._empty_state = tk.Frame(self._reader, bg=C['bg'])
        self._empty_state.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        center = tk.Frame(self._empty_state, bg=C['bg'])
        center.place(relx=0.5, rely=0.45, anchor='center')

        tk.Label(center, text="逐字稿", bg=C['bg'], fg=C['fg_dim'],
                 font=(UI_FONT, 22)).pack(pady=(0, 24))

        for key, desc in [(f'{MOD_KEY}+N', '新建文稿'),
                          (f'{MOD_KEY}+O', '导入文稿'),
                          (f'{MOD_KEY}+⇧+I', '导入 Word / PDF')]:
            row = tk.Frame(center, bg=C['bg'])
            row.pack(fill=tk.X, pady=5)
            tk.Label(row, text=key, bg=C['bg_input'], fg=C['fg_dim'],
                     font=(UI_FONT, 11), padx=10, pady=4,
                     highlightbackground=C['border'], highlightthickness=1
                     ).pack(side=tk.LEFT)
            tk.Label(row, text=f'  {desc}', bg=C['bg'], fg=C['fg_hint'],
                     font=(UI_FONT, 12)).pack(side=tk.LEFT)

        # text 初始不 pack，加载文件后再显示
        self.text.pack_forget()

        # 内容区限宽：reader 宽度变化时动态调整 text 的 padx 使正文列 ≤ 720px
        self._reader.bind('<Configure>', self._on_reader_resize)

    # ── 侧栏切换 ───────────────────────────────────────────

    def _on_text_scroll(self, first, last):
        """拦截滚动事件，同步更新滚动条和进度条"""
        self._vsb_text.set(first, last)
        self._update_progress(float(first), float(last))

    def _update_progress(self, first=None, last=None):
        if first is None:
            try:
                first, last = (float(v) for v in self.text.yview())
            except Exception:
                return
        pct = int(last * 100)
        if last >= 0.999:
            pct = 100
        # 进度百分比合并到顶栏 meta 标签
        self._update_meta_label(pct if self._current_file else None)
        try:
            total_w = self._progress_track.winfo_width()
            fill_w  = int(total_w * min(last, 1.0))
            self._progress_fill.place(x=0, y=0, relheight=1, width=fill_w)
        except Exception:
            pass

    def _update_meta_label(self, pct=None):
        """更新顶栏右侧 '字数 · 进度%' 标签"""
        chars_text = getattr(self, '_chars_text', '')
        if not self._current_file:
            self._lbl_meta.config(text='')
            return
        parts = []
        if chars_text:
            parts.append(chars_text)
        if pct is not None:
            parts.append(f'{pct}%')
        self._lbl_meta.config(text='  ·  '.join(parts))

    def _on_reader_resize(self, event):
        """动态调整 text padx，限制正文列宽 ≤ 720px，超宽时居中（防抖）"""
        if hasattr(self, '_resize_after') and self._resize_after:
            self.after_cancel(self._resize_after)
        self._resize_after = self.after(80, lambda w=event.width: self._apply_padx(w))

    def _apply_padx(self, reader_width):
        self._resize_after = None
        MAX_TEXT_W = 720
        MIN_PADX   = 64
        available  = reader_width - 8   # 减去滚动条宽度
        if available > MAX_TEXT_W + MIN_PADX * 2:
            padx = (available - MAX_TEXT_W) // 2
        else:
            padx = MIN_PADX
        new_padx = max(MIN_PADX, padx)
        # 仅在数值变化超过 4px 时才更新，避免频繁 config 干扰输入
        if abs(new_padx - getattr(self, '_current_padx', 0)) > 4:
            self._current_padx = new_padx
            try:
                self.text.config(padx=new_padx)
            except Exception:
                pass

    def _on_window_resize(self, event):
        """窗口变窄时自动折叠右侧批注面板"""
        if event.widget is not self:
            return
        if event.width < 1050 and getattr(self, '_annot_panel_visible', False):
            if self.annot_mgr:
                self.annot_mgr.toggle_panel()

    def _toggle_sidebar(self):
        if self._sidebar_visible:
            self._sidebar.pack_forget()
            self._sidebar_visible = False
        else:
            self._sidebar.pack(side=tk.LEFT, fill=tk.Y, before=self._reader)
            self._sidebar_visible = True

    # ── 作者设置 ───────────────────────────────────────────

    def _get_author(self):
        prefs = self.store.files.get('__prefs__', {})
        return prefs.get('author', '') if isinstance(prefs, dict) else ''

    def _set_author(self):
        current = self._get_author()
        self._show_entry(
            title='设置作者名',
            hint=current,
            cb=self._save_author,
        )

    def _save_author(self, name):
        prefs = self.store.files.get('__prefs__', {})
        if not isinstance(prefs, dict):
            prefs = {}
        prefs['author'] = name.strip()
        self.store.files['__prefs__'] = prefs
        self.store.save()
        if hasattr(self, '_refresh_author_lbl'):
            self._refresh_author_lbl()

    # ── 弹窗输入框 ─────────────────────────────────────────

    def _show_entry(self, hint='', cb=None, title=''):
        """弹出深色主题小对话框，获取用户输入后回调 cb(name)"""
        dlg = tk.Toplevel(self)
        dlg.title('')
        dlg.configure(bg=C['bg_input'])
        dlg.resizable(False, False)
        dlg.transient(self)
        dlg.grab_set()

        # 居中在主窗口
        self.update_idletasks()
        pw, ph = self.winfo_width(), self.winfo_height()
        px, py = self.winfo_rootx(), self.winfo_rooty()
        dw, dh = 300, 110
        dlg.geometry(f'{dw}x{dh}+{px + (pw - dw)//2}+{py + (ph - dh)//2}')

        # 标题行
        if title:
            tk.Label(dlg, text=title, bg=C['bg_input'], fg=C['fg_dim'],
                     font=('PingFang SC', 11)).pack(anchor='w', padx=16, pady=(14, 4))

        # 输入框
        evar = tk.StringVar(value=hint)
        entry = tk.Entry(dlg, textvariable=evar, bg=C['bg'],
                         fg=C['fg'], insertbackground=C['fg'],
                         font=('PingFang SC', 13), relief=tk.FLAT,
                         highlightbackground=C['accent_dim'],
                         highlightthickness=1)
        entry.pack(fill=tk.X, padx=16, ipady=6)
        entry.select_range(0, tk.END)
        entry.focus_force()

        # 确认 / 取消
        def ok(e=None):
            val = evar.get().strip()
            dlg.destroy()
            if val and cb:
                cb(val)

        def cancel(e=None):
            dlg.destroy()

        entry.bind('<Return>', ok)
        entry.bind('<Escape>', cancel)

        btn_row = tk.Frame(dlg, bg=C['bg_input'])
        btn_row.pack(fill=tk.X, padx=16, pady=(10, 0))

        for txt, fn, fg_normal in [('取消', cancel, C['fg_dim']),
                                    ('确定', ok, C['accent'])]:
            b = tk.Label(btn_row, text=txt, bg=C['bg_input'], fg=fg_normal,
                         font=('PingFang SC', 12), cursor='hand2', padx=10)
            b.pack(side=tk.RIGHT, padx=(6, 0))
            b.bind('<Button-1>', fn)
            b.bind('<Enter>', lambda e, w=b, c=fg_normal: w.config(fg=C['fg']))
            b.bind('<Leave>', lambda e, w=b, c=fg_normal: w.config(fg=c))

        dlg.wait_window()

    # ── 树形 ───────────────────────────────────────────────

    def _save_open_states(self):
        """把当前树中各 tag 节点的展开状态存入 store（始终用新 dict）"""
        states = {}
        def walk(node):
            v = self.tree.item(node, 'values')
            if v and v[1] == 'tag' and v[0] not in ('__untagged__', '__pin_hdr__'):
                states[v[0]] = bool(self.tree.item(node, 'open'))
            for child in self.tree.get_children(node):
                walk(child)
        for root in self.tree.get_children(''):
            walk(root)
        if states:  # 只在有内容时写入，防止覆盖为空
            self.store.files['__open__'] = states

    def _refresh_tree(self):
        # 只在树有内容时才保存（避免第一次调用时用空树覆盖已存状态）
        if self.tree.get_children(''):
            self._save_open_states()
        open_st = self.store.files.get('__open__', {})
        if not isinstance(open_st, dict):
            open_st = {}

        self.tree.delete(*self.tree.get_children())
        self._tooltip_data.clear()
        self._hide_tooltip()

        pinned = self.store.get_pinned()
        normal = [t for t in self.store.get_roots()
                  if not self.store.tags[t].get('pinned')]

        # ── 置顶标签区 ──
        if pinned:
            hdr = self.tree.insert('', 'end', text="  置顶标签",
                                   values=('__pin_hdr__', 'header'),
                                   open=True, tags=('header',))
            for tid in pinned:
                self._ins_tag(hdr, tid, 0, open_st)

        # ── 普通标签区 ──
        for tid in normal:
            self._ins_tag('', tid, 0, open_st)

        # ── 未分类 ──
        untagged = [fp for fp, tags in self.store.files.items()
                    if not fp.startswith('__') and isinstance(tags, list) and not tags]
        if untagged:
            node = self.tree.insert('', 'end', text="  未分类",
                                    values=('__untagged__', 'tag'), open=True)
            for fp in untagged:
                name = os.path.basename(fp)
                iid = self.tree.insert(node, 'end',
                                 text=self._trim_name(name),
                                 values=(fp, 'file'), tags=('file',))
                self._tooltip_data[iid] = name

    def _trim_name(self, name, limit=18):
        """截断过长文件名，保留后缀"""
        if len(name) <= limit:
            return name
        dot = name.rfind('.')
        if dot > 0 and len(name) - dot <= 8:
            ext = name[dot:]
            stem = name[:dot]
            keep = limit - len(ext) - 1
            return stem[:keep] + '…' + ext
        return name[:limit - 1] + '…'

    def _on_tree_motion(self, event):
        iid = self.tree.identify_row(event.y)
        if iid == self._tooltip_item:
            # 同一行：只更新 tooltip 位置
            if self._tooltip_win:
                x = event.x_root + 14
                y = event.y_root + 14
                self._tooltip_win.wm_geometry(f'+{x}+{y}')
            return
        self._hide_tooltip()
        self._tooltip_item = iid
        full = self._tooltip_data.get(iid)
        if not full:
            return
        # 只有截断了才弹 tooltip
        if self._trim_name(full) == full:
            return
        win = tk.Toplevel(self)
        win.wm_overrideredirect(True)
        win.wm_attributes('-topmost', True)
        tk.Label(win, text=full,
                 bg='#1e1e2e', fg='#d4d4e8',
                 font=('PingFang SC', 12),
                 padx=10, pady=5,
                 relief=tk.FLAT,
                 bd=0).pack()
        x = event.x_root + 14
        y = event.y_root + 14
        win.wm_geometry(f'+{x}+{y}')
        self._tooltip_win = win

    def _on_toc_motion(self, event):
        iid = self._toc_tree.identify_row(event.y)
        if iid == self._tooltip_item:
            if self._tooltip_win:
                self._tooltip_win.wm_geometry(f'+{event.x_root+14}+{event.y_root+14}')
            return
        self._hide_tooltip()
        self._tooltip_item = iid
        if not iid:
            return
        full = self._toc_tree.item(iid, 'text').strip()
        if not full or full == '暂无目录条目':
            return
        win = tk.Toplevel(self)
        win.wm_overrideredirect(True)
        win.wm_attributes('-topmost', True)
        tk.Label(win, text=full,
                 bg='#1e1e2e', fg='#d4d4e8',
                 font=('PingFang SC', 12),
                 padx=10, pady=5,
                 relief=tk.FLAT, bd=0).pack()
        win.wm_geometry(f'+{event.x_root+14}+{event.y_root+14}')
        self._tooltip_win = win

    def _hide_tooltip(self, event=None):
        if self._tooltip_win:
            self._tooltip_win.destroy()
            self._tooltip_win = None
        self._tooltip_item = None

    def _ins_tag(self, parent, tid, depth, open_st=None):
        tag = self.store.tags[tid]
        n   = self.store.count_under(tid)
        cnt = f"  {n}" if n else ""
        pin = "  📌" if tag.get('pinned') and depth == 0 else ""
        lbl = f"  {tag['name']}{cnt}{pin}"   # Flomo 风格：无 # 前缀，靠缩进表达层级
        # 默认展开；若有保存状态则恢复
        is_open = open_st.get(tid, True) if open_st else True
        node = self.tree.insert(parent, 'end', text=lbl,
                                values=(tid, 'tag'), open=is_open, tags=('tag',))
        for cid in tag['children']:
            self._ins_tag(node, cid, depth + 1, open_st)
        for fp in self.store.files_for(tid):
            name = os.path.basename(fp)
            iid = self.tree.insert(node, 'end',
                             text=self._trim_name(name),
                             values=(fp, 'file'), tags=('file',))
            self._tooltip_data[iid] = name

    # ── 事件 ───────────────────────────────────────────────

    def _on_tree_toggle(self, _e):
        """折叠/展开后延一帧保存，确保 open 状态已更新"""
        self.after(50, self._do_save_open)

    def _do_save_open(self):
        self._save_open_states()
        self.store.save()

    def _on_tree_select(self, _e):
        sel = self.tree.selection()
        if not sel:
            return
        v = self.tree.item(sel[0], 'values')
        if v and v[1] == 'header':
            self.tree.selection_remove(sel[0])
            return
        if v and v[1] == 'file':
            self._load(v[0])

    def _on_rclick(self, event):
        item = self.tree.identify_row(event.y)
        self._ctx.delete(0, tk.END)
        if not item:
            self._ctx.add_command(label="新建文稿…",        command=self._new_doc)
            self._ctx.add_separator()
            self._ctx.add_command(label="导入 TXT…",        command=self._import)
            self._ctx.add_command(label="导入 Word…",       command=self._import_word)
            self._ctx.add_command(label="导入 PDF…",        command=self._import_pdf)
            self._ctx.add_separator()
            self._ctx.add_command(label="导入工程文件…",    command=self._import_zzg)
            self._ctx.add_separator()
            self._ctx.add_command(label="新建标签…",        command=self._new_root_tag)
        else:
            self.tree.selection_set(item)
            v    = self.tree.item(item, 'values')
            kind = v[1] if v else None
            if kind == 'tag' and v[0] not in ('__untagged__', '__pin_hdr__'):
                tid = v[0]
                self._ctx.add_command(label="新建标签…",
                    command=self._new_root_tag)
                self._ctx.add_command(label="新建子标签…",
                    command=lambda: self._new_child(tid))
                self._ctx.add_separator()
                self._ctx.add_command(label="重命名",
                    command=lambda: self._rename(tid))
                # 置顶 / 取消置顶（仅顶级标签）
                tag = self.store.tags.get(tid, {})
                if not tag.get('parent'):
                    if tag.get('pinned'):
                        self._ctx.add_command(label="取消置顶",
                            command=lambda t=tid: (self.store.pin_tag(t, False),
                                                   self._refresh_tree()))
                    else:
                        self._ctx.add_command(label="置顶",
                            command=lambda t=tid: (self.store.pin_tag(t, True),
                                                   self._refresh_tree()))
                self._ctx.add_separator()
                self._ctx.add_command(label="导入 TXT 到此标签…",
                    command=lambda: self._import(tid))
                self._ctx.add_command(label="导入 Word 到此标签…",
                    command=lambda t=tid: self._import_word(t))
                self._ctx.add_command(label="导入 PDF 到此标签…",
                    command=lambda t=tid: self._import_pdf(t))
                self._ctx.add_separator()
                self._ctx.add_command(label="删除标签",
                    command=lambda: self._delete(tid))
            elif kind == 'file':
                fp   = v[0]
                pv   = self.tree.item(self.tree.parent(item), 'values')
                ptid = pv[0] if pv and pv[1] == 'tag' and pv[0] != '__untagged__' else None
                # 移动到标签子菜单
                all_tags = self._collect_all_tags()
                if all_tags:
                    move_menu = tk.Menu(self._ctx, tearoff=0)
                    for tid_m, tname_m in all_tags:
                        if tid_m != ptid:
                            move_menu.add_command(
                                label=tname_m,
                                command=lambda f=fp, t=tid_m: self._move_file_to_tag(f, t))
                    self._ctx.add_cascade(label="移动到标签", menu=move_menu)
                    self._ctx.add_separator()
                if ptid:
                    self._ctx.add_command(label="从此标签移除",
                        command=lambda f=fp, t=ptid: (self.store.remove_file(f, t),
                                                       self._refresh_tree()))
                    self._ctx.add_separator()
                self._ctx.add_command(label="从侧栏删除",
                    command=lambda f=fp: self._remove_file_from_store(f))
                self._ctx.add_separator()
                self._ctx.add_command(label="在 Finder 中显示",
                    command=lambda f=fp: self._reveal_in_finder(f))
                self._ctx.add_separator()
                self._ctx.add_command(label="导出为 PDF…",
                    command=self._export_pdf)
                self._ctx.add_command(label="导出为 Word…",
                    command=self._export_word)
                self._ctx.add_separator()
                self._ctx.add_command(label="导出工程文件…",
                    command=self._export_zzg)
        self._ctx.tk_popup(event.x_root, event.y_root)

    def _collect_all_tags(self):
        """返回所有标签列表 [(tid, 名称路径), ...]，含层级缩进"""
        result = []
        def walk(tid, depth):
            tag = self.store.tags.get(tid)
            if not tag:
                return
            prefix = '  ' * depth
            result.append((tid, f"{prefix}{tag['name']}"))
            for cid in tag.get('children', []):
                walk(cid, depth + 1)
        for root_tid in self.store.get_roots():
            walk(root_tid, 0)
        return result

    def _move_file_to_tag(self, fp, tag_id):
        """将文件移动到指定标签（保留其他标签关联）"""
        self.store.add_file(fp, tag_id)
        self._refresh_tree()

    def _reveal_in_finder(self, fp):
        import subprocess
        if os.path.exists(fp):
            subprocess.run(['open', '-R', fp])
        else:
            messagebox.showwarning("提示", f"原文件不存在：\n{fp}")

    def _remove_file_from_store(self, fp):
        name = os.path.basename(fp)
        if messagebox.askyesno("删除", f"从侧栏移除「{name}」？\n（原文件不受影响）"):
            if fp in self.store.files:
                del self.store.files[fp]
                self.store.save()
            if self._current_file == fp:
                self._current_file = None
                self.text.delete('1.0', tk.END)
                self._lbl_file.config(text="逐字稿")
                self._chars_text = ''
                self._update_meta_label()
            self._refresh_tree()

    def _new_root_tag(self):
        self._show_entry(title='新建标签',
                         cb=lambda n: (self.store.add_tag(n), self._refresh_tree()))

    def _new_child(self, ptid):
        self._show_entry(title='新建子标签',
                         cb=lambda n: (self.store.add_tag(n, parent_id=ptid),
                                       self._refresh_tree()))

    def _rename(self, tid):
        self._show_entry(title='重命名',
                         hint=self.store.tags[tid]['name'],
                         cb=lambda n: (self.store.rename_tag(tid, n),
                                       self._refresh_tree()))

    def _delete(self, tid):
        name = self.store.tags[tid]['name']
        if messagebox.askyesno("删除", f"确定删除「{name}」？"):
            self.store.delete_tag(tid)
            self._refresh_tree()

    def _import(self, tag_id=None):
        if tag_id is None:
            sel = self.tree.selection()
            if sel:
                v = self.tree.item(sel[0], 'values')
                if v and v[1] == 'tag' and v[0] != '__untagged__':
                    tag_id = v[0]

        r = subprocess.run(
            ['osascript', '-e',
             'set fs to choose file with prompt "选择文稿" with multiple selections allowed\n'
             'set out to ""\n'
             'repeat with f in fs\n'
             'set out to out & POSIX path of f & "\n"\n'
             'end repeat\n'
             'out'],
            capture_output=True, text=True)

        paths = [p.strip() for p in r.stdout.splitlines() if p.strip()]
        if not paths:
            return

        for p in paths:
            self.store.add_file(p, tag_id)
        self._refresh_tree()
        self._load(paths[-1])

    # ── 新建文稿 ───────────────────────────────────────────────

    def _new_doc(self):
        path = filedialog.asksaveasfilename(
            title="新建文稿",
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt")],
        )
        if not path:
            return
        open(path, 'w', encoding='utf-8').write('')
        self.store.add_file(path, None)
        self._refresh_tree()
        self._load(path)

    # ── 导入 Word ──────────────────────────────────────────────

    def _import_word(self, tag_id=None):
        path = filedialog.askopenfilename(
            title="导入 Word 文件",
            filetypes=[("Word 文件", "*.docx *.doc")],
        )
        if not path:
            return
        try:
            import docx
            doc  = docx.Document(path)
            text = '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            messagebox.showerror("导入失败", str(e))
            return
        self.store.add_file(path, tag_id)
        self._refresh_tree()
        self._load_text_content(path, text)

    # ── 导入 PDF ───────────────────────────────────────────────

    def _import_pdf(self, tag_id=None):
        path = filedialog.askopenfilename(
            title="导入 PDF 文件",
            filetypes=[("PDF 文件", "*.pdf")],
        )
        if not path:
            return
        try:
            import fitz  # PyMuPDF
            doc  = fitz.open(path)
            text = '\n\n'.join(page.get_text() for page in doc)
            doc.close()
        except Exception as e:
            messagebox.showerror("导入失败", str(e))
            return
        self.store.add_file(path, tag_id)
        self._refresh_tree()
        self._load_text_content(path, text)

    # ── 导出 PDF ───────────────────────────────────────────────

    def _export_pdf(self):
        if not self._current_file:
            messagebox.showwarning("提示", "请先打开文稿")
            return
        default = os.path.splitext(self._current_file)[0] + '.pdf'
        path = filedialog.asksaveasfilename(
            title="导出为 PDF",
            defaultextension=".pdf",
            initialfile=os.path.basename(default),
            filetypes=[("PDF 文件", "*.pdf")],
        )
        if not path:
            return
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 注册 LXGW WenKai 字体（支持中文）
            font_path = os.path.join(FONTS_DIR, 'LXGWWenKai-Regular.ttf')
            pdfmetrics.registerFont(TTFont('WenKai', font_path))

            content = self.text.get('1.0', tk.END)
            W, H    = A4
            margin  = 60
            font_sz = 12
            line_h  = font_sz * 1.8
            max_w   = W - margin * 2

            c     = rl_canvas.Canvas(path, pagesize=A4)
            c.setFont('WenKai', font_sz)
            y     = H - margin

            for para in content.split('\n'):
                # 简单分词折行
                words = list(para) if para else ['']
                line  = ''
                for ch in words:
                    test = line + ch
                    if c.stringWidth(test, 'WenKai', font_sz) > max_w:
                        c.drawString(margin, y, line)
                        y -= line_h
                        if y < margin:
                            c.showPage()
                            c.setFont('WenKai', font_sz)
                            y = H - margin
                        line = ch
                    else:
                        line = test
                c.drawString(margin, y, line)
                y -= line_h * 1.2  # 段落间距
                if y < margin:
                    c.showPage()
                    c.setFont('WenKai', font_sz)
                    y = H - margin

            c.save()
            messagebox.showinfo("导出成功", f"已保存到：\n{path}")
        except Exception as e:
            messagebox.showerror("导出失败", str(e))

    # ── 导出 Word ──────────────────────────────────────────────

    def _export_word(self):
        if not self._current_file:
            messagebox.showwarning("提示", "请先打开文稿")
            return
        default = os.path.splitext(self._current_file)[0] + '.docx'
        path = filedialog.asksaveasfilename(
            title="导出为 Word",
            defaultextension=".docx",
            initialfile=os.path.basename(default),
            filetypes=[("Word 文件", "*.docx")],
        )
        if not path:
            return
        try:
            import docx
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn

            doc  = docx.Document()
            # 页边距
            for sec in doc.sections:
                sec.top_margin    = Cm(2.5)
                sec.bottom_margin = Cm(2.5)
                sec.left_margin   = Cm(3)
                sec.right_margin  = Cm(3)

            content = self.text.get('1.0', tk.END)
            for line in content.split('\n'):
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(12)

            doc.save(path)
            messagebox.showinfo("导出成功", f"已保存到：\n{path}")
        except Exception as e:
            messagebox.showerror("导出失败", str(e))

    # ── 工程文件导出/导入（.zzg）─────────────────────────────────

    def _export_annots_md(self):
        """将当前文件的批注导出为 Markdown 文件"""
        if not self._current_file or not self.annot_mgr:
            messagebox.showwarning("提示", "请先打开文稿")
            return
        annots = self.annot_mgr.store.get_for_file(self._current_file)
        if not annots:
            messagebox.showinfo("提示", "当前文稿暂无批注")
            return
        stem = os.path.splitext(os.path.basename(self._current_file))[0]
        path = filedialog.asksaveasfilename(
            title="导出批注为 Markdown",
            defaultextension=".md",
            initialfile=stem + '_批注.md',
            filetypes=[("Markdown 文件", "*.md"), ("文本文件", "*.txt")],
        )
        if not path:
            return
        from annotation_manager import ANNOT_STYLES as _AS
        lines = [f"# {stem} — 批注导出\n"]
        for a in annots:
            style = _AS.get(a['type'], {})
            label = style.get('label', a['type'])
            author = a.get('author', '')
            date   = a.get('created_at', '')[:10]
            meta   = '  '.join(filter(None, [f'@{author}' if author else '', date]))
            lines.append(f"## [{label}]")
            lines.append(f"> {a['text'].replace(chr(10), ' ')}")
            note = a.get('note', '').strip()
            if note:
                lines.append(f"\n{note}")
            if meta:
                lines.append(f"\n_{meta}_")
            lines.append("\n---\n")
        try:
            open(path, 'w', encoding='utf-8').write('\n'.join(lines))
            messagebox.showinfo("导出成功", f"批注已导出：\n{path}")
        except Exception as e:
            messagebox.showerror("导出失败", str(e))

    def _export_zzg(self):
        """把当前文稿 + 全部标注打包成 .zzg 工程文件"""
        if not self._current_file:
            messagebox.showwarning("提示", "请先打开文稿")
            return
        import zipfile, json as _json
        stem = os.path.splitext(os.path.basename(self._current_file))[0]
        path = filedialog.asksaveasfilename(
            title="导出工程文件",
            defaultextension=".zzg",
            initialfile=stem + '.zzg',
            filetypes=[("逐字稿工程", "*.zzg")],
        )
        if not path:
            return
        try:
            content = self.text.get('1.0', tk.END)
            annots  = self.store.files.get('__annotations__', {})
            file_annots = annots.get(self._current_file, [])
            toc_entries = self._get_toc_entries(self._current_file)
            meta = {
                'version':   2,
                'filename':  os.path.basename(self._current_file),
                'ext':       os.path.splitext(self._current_file)[1].lower(),
            }
            with zipfile.ZipFile(path, 'w', zipfile.ZIP_LZMA) as zf:
                zf.writestr('meta.json',        _json.dumps(meta, ensure_ascii=False))
                zf.writestr('content.txt',      content.encode('utf-8'))
                zf.writestr('annotations.json', _json.dumps(file_annots, ensure_ascii=False, indent=2))
                zf.writestr('toc.json',         _json.dumps(toc_entries, ensure_ascii=False, indent=2))
            messagebox.showinfo("导出成功", f"工程文件已保存：\n{path}")
        except Exception as e:
            messagebox.showerror("导出失败", str(e))

    def _import_zzg(self):
        """导入 .zzg 工程文件，还原文稿内容与全部标注"""
        import zipfile, json as _json
        path = filedialog.askopenfilename(
            title="导入工程文件",
            filetypes=[("逐字稿工程", "*.zzg")],
        )
        if not path:
            return
        try:
            with zipfile.ZipFile(path, 'r') as zf:
                meta        = _json.loads(zf.read('meta.json').decode('utf-8'))
                content     = zf.read('content.txt').decode('utf-8')
                file_annots = _json.loads(zf.read('annotations.json').decode('utf-8'))
                # toc.json 兼容旧版（version 1）
                if 'toc.json' in zf.namelist():
                    toc_entries = _json.loads(zf.read('toc.json').decode('utf-8'))
                else:
                    toc_entries = []

            # 以工程文件路径作为虚拟文件路径（保证唯一）
            vpath = path

            # 写入标注
            all_annots = self.store.files.get('__annotations__', {})
            if not isinstance(all_annots, dict):
                all_annots = {}
            for a in file_annots:
                a['file'] = vpath
            all_annots[vpath] = file_annots
            self.store.files['__annotations__'] = all_annots

            # 写入目录
            self._save_toc_entries(vpath, toc_entries)

            # 注册到侧栏
            self.store.add_file(vpath, None)
            self._refresh_tree()

            # 加载内容（只读保护 pdf/docx，txt 可编辑）
            ext = meta.get('ext', '.txt')
            self.text.config(state=tk.NORMAL)
            self._load_text_content(vpath, content)
            if ext in ('.pdf', '.docx', '.doc'):
                self.text.config(state=tk.DISABLED)

            messagebox.showinfo("导入成功",
                f"已导入：{meta.get('filename','')}\n"
                f"批注数：{len(file_annots)}　目录条目：{len(toc_entries)}")
        except Exception as e:
            messagebox.showerror("导入失败", str(e))

    def _font_change(self, delta):
        self._font_size = max(12, min(36, self._font_size + delta))
        self.text.config(font=(self._font_family, self._font_size))
        self._save_prefs()

    def _save_prefs(self):
        prefs = self.store.files.get('__prefs__', {})
        if not isinstance(prefs, dict):
            prefs = {}
        prefs['font_size'] = self._font_size
        prefs['font_family'] = self._font_family
        prefs['line_spacing'] = self._line_spacing
        self.store.files['__prefs__'] = prefs
        self.store.save()

    def _spacing_change(self, delta):
        self._line_spacing = max(4, min(40, self._line_spacing + delta))
        self.text.config(spacing2=self._line_spacing, spacing3=self._line_spacing)
        self._save_prefs()

    def _apply_font(self, family):
        self._font_family = family
        self.text.config(font=(self._font_family, self._font_size))
        self._save_prefs()

    def _show_font_picker(self):
        win = tk.Toplevel(self)
        win.title("选择字体")
        win.geometry("460x520")
        win.configure(bg=C['bg'])
        win.resizable(False, False)
        win.transient(self)
        win.grab_set()

        tk.Label(win, text="选择字体", bg=C['bg'], fg=C['fg'],
                 font=('PingFang SC', 13, 'bold')).pack(pady=(20, 8))

        # 搜索框
        search_var = tk.StringVar()
        search_entry = tk.Entry(win, textvariable=search_var,
                                bg=C['bg_input'], fg=C['fg'],
                                insertbackground=C['fg'],
                                font=('PingFang SC', 12),
                                relief=tk.FLAT, bd=8)
        search_entry.pack(fill=tk.X, padx=20)
        tk.Frame(win, bg=C['border'], height=1).pack(fill=tk.X, padx=20)

        # 字体列表
        list_frame = tk.Frame(win, bg=C['bg'])
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=8)

        sb = ttk.Scrollbar(list_frame, style='Thin.Vertical.TScrollbar')
        lb = tk.Listbox(list_frame, bg=C['bg_input'], fg=C['fg_tag'],
                        selectbackground=C['accent_dim'],
                        selectforeground=C['fg'],
                        font=('PingFang SC', 12),
                        relief=tk.FLAT, bd=0,
                        yscrollcommand=sb.set,
                        activestyle='none')
        sb.config(command=lb.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 内置字体置顶（显示名 → tkinter字族名），其余系统字体按中文优先
        builtin_display = [(disp, tk_name) for disp, tk_name in BUILTIN_FONTS.values()]
        builtin_tk_names = {tk_name for _, tk_name in builtin_display}
        system_fonts = sorted(
            [(f, f) for f in tkfont.families() if f not in builtin_tk_names],
            key=lambda x: (0 if any(c > '\u4e00' for c in x[0]) else 1, x[0].lower()))
        all_fonts = builtin_display + system_fonts  # [(显示名, tk字族名), ...]

        filtered = []  # 当前显示的 [(显示名, tk字族名)]

        def refresh_list(*_):
            nonlocal filtered
            kw = search_var.get().lower()
            filtered = [(d, t) for d, t in all_fonts if kw in d.lower() or kw in t.lower()]
            lb.delete(0, tk.END)
            for disp, _ in filtered:
                lb.insert(tk.END, disp)
            for i, (_, tk_name) in enumerate(filtered):
                if tk_name == self._font_family:
                    lb.selection_set(i)
                    lb.see(i)
                    break

        search_var.trace_add('write', refresh_list)
        refresh_list()

        # 预览区
        preview_label = tk.Label(win, text="预览：天地玄黄 宇宙洪荒",
                                  bg=C['bg_input'], fg=C['fg'],
                                  font=(self._font_family, 16),
                                  pady=12, relief=tk.FLAT)
        preview_label.pack(fill=tk.X, padx=20, pady=(0, 8))

        def on_select(_e=None):
            sel = lb.curselection()
            if sel and sel[0] < len(filtered):
                tk_name = filtered[sel[0]][1]
                try:
                    preview_label.config(font=(tk_name, 16))
                except Exception:
                    pass

        lb.bind('<<ListboxSelect>>', on_select)

        # 按钮
        btn_frame = tk.Frame(win, bg=C['bg'])
        btn_frame.pack(fill=tk.X, padx=20, pady=(0, 16))

        def confirm():
            sel = lb.curselection()
            if sel and sel[0] < len(filtered):
                self._apply_font(filtered[sel[0]][1])
            win.destroy()

        tk.Button(btn_frame, text="应用", bg=C['accent'], fg='white',
                  font=('PingFang SC', 12), relief=tk.FLAT, bd=0,
                  padx=20, pady=6, cursor='hand2', command=confirm,
                  activebackground=C['accent_dim'], activeforeground='white'
                  ).pack(side=tk.RIGHT, padx=(8, 0))
        tk.Button(btn_frame, text="取消", bg=C['bg_input'], fg=C['fg_file'],
                  font=('PingFang SC', 12), relief=tk.FLAT, bd=0,
                  padx=20, pady=6, cursor='hand2', command=win.destroy,
                  activebackground=C['bg_input'], activeforeground=C['fg']
                  ).pack(side=tk.RIGHT)

        search_entry.focus_set()
        win.bind('<Return>', lambda e: confirm())
        win.bind('<Escape>', lambda e: win.destroy())

    def _import_font_file(self):
        r = subprocess.run(
            ['osascript', '-e',
             'set f to choose file with prompt "选择字体文件（.ttf 或 .otf）" '
             'of type {"public.truetype-ttf-font", "public.opentype-font", '
             '"com.adobe.postscript-font", "org.gnu.gnu-type-1-font"}\n'
             'POSIX path of f'],
            capture_output=True, text=True)
        path = r.stdout.strip()
        if not path:
            return
        font_dir = os.path.expanduser('~/Library/Fonts')
        dest = os.path.join(font_dir, os.path.basename(path))
        try:
            shutil.copy2(path, dest)
        except Exception as e:
            messagebox.showerror("导入失败", str(e))
            return
        messagebox.showinfo("导入成功",
            f"字体已安装到系统。\n\n请前往「显示 → 选择字体」搜索并应用新字体。")
        # 重新打开字体选择器（新字体可能需要重启才能在列表出现）
        self._show_font_picker()

    def _save_position(self):
        if not self._current_file:
            return
        try:
            idx    = self.text.index('@73,41')
            offset = int(self.text.count('1.0', idx, 'chars')[0])
            # 分段模式：加上当前段的起始偏移，保存全文绝对位置
        except Exception:
            offset = 0
        positions = self.store.files.get('__pos__', {})
        if not isinstance(positions, dict):
            positions = {}
        positions[self._current_file] = offset
        self.store.files['__pos__'] = positions
        self.store.save()

    def _load_text_content(self, path, content):
        """直接加载已提取的文字内容（Word/PDF 用，不写磁盘）"""
        self._flush_reading_time()
        import time
        self._reading_start = time.time()
        self._save_position()

        self._current_file = path
        self._clear_search_highlights()
        self._search_matches = []
        self._search_idx     = -1
        self._search_count_lbl.config(text="")
        self.text.config(state=tk.NORMAL)
        self.text.delete('1.0', tk.END)
        self.text.insert('1.0', content)
        self.text.config(state=tk.DISABLED)  # PDF/Word 锁定为只读

        zh = len(re.findall(r'[\u4e00-\u9fff]', content))
        self._lbl_file.config(text=f"{os.path.basename(path)}  🔒")
        self._chars_text = f"{zh:,} 字"
        self._update_meta_label()
        self._update_time_label()

        # 恢复标注和目录高亮（延后执行，先显示内容）
        _path = path
        def _deferred_restore2():
            if self.annot_mgr and self._current_file == _path:
                self.annot_mgr.on_file_loaded(_path)
            self._configure_heading_tags()
            if self._current_file == _path:
                self._restore_heading_tags()
            if self._sb_tab.get() == 'toc' and self._current_file == _path:
                self._refresh_toc()
        self.after(80, _deferred_restore2)
        self.lift()
        self.focus_force()

        pos_store    = self.store.files.get('__pos__', {})
        saved_val    = pos_store.get(path, 0) if isinstance(pos_store, dict) else 0
        saved_offset = 0 if isinstance(saved_val, float) else max(0, int(saved_val))

        def restore():
            try:
                self.text.update_idletasks()
                target = f'1.0+{saved_offset}c'
                self.text.see(target)
                self.text.update_idletasks()
                self.text.yview(target)
            except Exception:
                pass

        if saved_offset > 0:
            restore()
        self.after(250, restore)

    def _load(self, path):
        # 切换文件前：立即保存当前 txt 的未保存改动
        if self._current_file:
            prev_ext = os.path.splitext(self._current_file)[1].lower()
            if prev_ext not in ('.pdf', '.docx', '.doc', '.zzg') and self.text.edit_modified():
                try:
                    open(self._current_file, 'w', encoding='utf-8').write(
                        self.text.get('1.0', tk.END))
                except Exception:
                    pass
                if self._autosave_after:
                    self.after_cancel(self._autosave_after)
                    self._autosave_after = None

        # 结算上一个文件的阅读时长
        self._flush_reading_time()
        import time
        self._reading_start = time.time()  # 新文件立刻开始计时
        self._save_position()

        path = path.strip()
        ext = os.path.splitext(path)[1].lower()
        content = None

        if ext == '.zzg':
            try:
                import zipfile, json as _json
                with zipfile.ZipFile(path, 'r') as zf:
                    content = zf.read('content.txt').decode('utf-8')
                    meta    = _json.loads(zf.read('meta.json').decode('utf-8'))
                    # 若标注尚未写入 store（重启后重新点击），则恢复
                    all_annots = self.store.files.get('__annotations__', {})
                    if path not in all_annots:
                        file_annots = _json.loads(zf.read('annotations.json').decode('utf-8'))
                        for a in file_annots:
                            a['file'] = path
                        all_annots[path] = file_annots
                        self.store.files['__annotations__'] = all_annots
                orig_ext = meta.get('ext', '.txt')
            except Exception as e:
                messagebox.showerror("错误", f"无法读取工程文件：{e}")
                return
            self._load_text_content(path, content)
            if orig_ext not in ('.pdf', '.docx', '.doc'):
                self.text.config(state=tk.NORMAL)
            return
        elif ext == '.pdf':
            try:
                import fitz
                doc = fitz.open(path)
                content = '\n\n'.join(page.get_text() for page in doc)
                doc.close()
                # 提取 PDF 书签目录（首次加载）
                if path not in self.store.files.get('__toc__', {}):
                    self._extract_pdf_toc(path)
            except Exception as e:
                messagebox.showerror("错误", f"无法读取：{path}\n{e}")
                return
        elif ext in ('.docx', '.doc'):
            try:
                import docx
                doc = docx.Document(path)
                content = '\n'.join(p.text for p in doc.paragraphs)
                # 提取 Word 标题目录（首次加载）
                if path not in self.store.files.get('__toc__', {}):
                    self._extract_word_toc(doc, path)
            except Exception as e:
                messagebox.showerror("错误", f"无法读取：{path}\n{e}")
                return
        else:
            for enc in ('utf-8', 'gbk', 'utf-16'):
                try:
                    content = open(path, encoding=enc).read()
                    break
                except Exception:
                    continue
            if content is None:
                messagebox.showerror("错误", f"无法读取：{path}")
                return

        self._current_file = path
        # 记住最后打开的文件，下次启动时恢复
        self.store.files['__last_file__'] = path
        self._clear_search_highlights()
        self._search_matches = []
        self._search_idx = -1
        self._search_count_lbl.config(text="")
        # 先取出目标位置，插入前就准备好
        pos_store    = self.store.files.get('__pos__', {})
        saved_val    = pos_store.get(path, 0) if isinstance(pos_store, dict) else 0
        saved_offset = 0 if isinstance(saved_val, float) else max(0, int(saved_val))

        # 首次加载文件：隐藏空状态，显示文本区
        # 用 winfo_manager() 而非 winfo_ismapped()，后者在窗口 withdraw 时返回 False
        if hasattr(self, '_empty_state') and self._empty_state.winfo_manager():
            self._empty_state.pack_forget()
            self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        ext = os.path.splitext(path)[1].lower()
        self.text.config(state=tk.NORMAL, undo=False)
        self.text.delete('1.0', tk.END)
        self.text.insert('1.0', content)
        if ext in ('.pdf', '.docx', '.doc'):
            self.text.config(state=tk.DISABLED, cursor='arrow', undo=False)
        else:
            self.text.config(state=tk.NORMAL, cursor='xterm', undo=True)
        self.text.edit_reset()

        fname = os.path.basename(path)
        self._lbl_file.config(text=f"{fname}  🔒" if ext in ('.pdf', '.docx', '.doc') else fname)
        total = len(content)
        zh = sum(1 for c in content if '\u4e00' <= c <= '\u9fff')
        self._chars_text = f"{zh:,} 字" if zh > total * 0.3 else f"{total:,} 字"
        self._update_meta_label()
        self._update_time_label()

        # 恢复标注和目录高亮（延后执行，先让内容显示出来，避免卡顿感）
        _path = path
        def _deferred_restore():
            if self.annot_mgr and self._current_file == _path:
                self.annot_mgr.on_file_loaded(_path)
            if self._current_file == _path:
                self._restore_heading_tags()
            if self._sb_tab.get() == 'toc' and self._current_file == _path:
                self._refresh_toc()
        self.after(80, _deferred_restore)

        self.lift()
        self.focus_force()
        # txt 文件直接聚焦文本区，让用户可以立即输入
        if ext not in ('.pdf', '.docx', '.doc'):
            self.text.focus_set()

        # ── 立刻定位（在 focus_force 之后，避免被覆盖）──
        def restore():
            try:
                self.text.update_idletasks()
                target = f'1.0+{saved_offset}c'
                self.text.see(target)
                self.text.update_idletasks()
                self.text.yview(target)
            except Exception:
                pass

        if saved_offset > 0:
            restore()                 # 同步立刻执行，消除闪跳
        self.after(250, restore)      # 保底一次，确保 macOS 渲染完成后准确

    def _save(self):
        if not self._current_file:
            return
        ext = os.path.splitext(self._current_file)[1].lower()
        if ext in ('.pdf', '.docx', '.doc'):
            # PDF/Word 为只读格式，提示另存为 txt
            default = os.path.splitext(self._current_file)[0] + '.txt'
            path = filedialog.asksaveasfilename(
                title="另存为文本文件",
                defaultextension='.txt',
                initialfile=os.path.basename(default),
                filetypes=[("文本文件", "*.txt")],
            )
            if not path:
                return
            try:
                open(path, 'w', encoding='utf-8').write(self.text.get('1.0', tk.END))
                self.store.add_file(path, None)
                self._refresh_tree()
                self._current_file = path
                self._lbl_file.config(text=f"{os.path.basename(path)} ✓")
                self.after(2000, lambda: self._lbl_file.config(text=os.path.basename(path)))
            except Exception as e:
                messagebox.showerror("保存失败", str(e))
            return
        content = self.text.get('1.0', tk.END)
        try:
            open(self._current_file, 'w', encoding='utf-8').write(content)
            self._lbl_file.config(text=f"{os.path.basename(self._current_file)} ✓")
            self.after(2000, lambda: self._lbl_file.config(
                text=os.path.basename(self._current_file)))
        except Exception as e:
            messagebox.showerror("保存失败", str(e))

    # ── 自动保存 / 实时字数 ─────────────────────────────────

    def _on_key_release(self, event=None):
        """按键松开时更新字数并触发防抖自动保存（不碰 edit_modified）"""
        if not self._current_file:
            return
        ext = os.path.splitext(self._current_file)[1].lower()
        if ext in ('.pdf', '.docx', '.doc'):
            return


        # 字数统计防抖（1秒后执行，不在每次按键时阻塞）
        if hasattr(self, '_wordcount_after') and self._wordcount_after:
            self.after_cancel(self._wordcount_after)
        self._wordcount_after = self.after(1000, self._update_wordcount)

        # 自动保存（防抖 2 秒）
        if self._autosave_after:
            self.after_cancel(self._autosave_after)
        self._autosave_after = self.after(2000, self._autosave)

    def _update_wordcount(self):
        """防抖字数统计，避免大文件打字卡顿"""
        self._wordcount_after = None
        if not self._current_file:
            return
        try:
            content = self.text.get('1.0', tk.END)
            total = len(content.strip())
            zh = sum(1 for c in content if '\u4e00' <= c <= '\u9fff')
            self._chars_text = f"{zh:,} 字" if zh > total * 0.3 else f"{total:,} 字"
            self._update_meta_label()
        except Exception:
            pass

    def _check_notion_heading(self, event=None):
        """Notion 风格：行首 # 空格 标题 + Enter，自动去掉前缀并标记目录"""
        if not self._current_file:
            return
        ext = os.path.splitext(self._current_file)[1].lower()
        if ext in ('.pdf', '.docx', '.doc'):
            return
        try:
            insert = self.text.index(tk.INSERT)
            line_no = insert.split('.')[0]
            line_text = self.text.get(f'{line_no}.0', f'{line_no}.end')
            m = re.match(r'^(#{1,7}) (.+)', line_text)
            if not m:
                return   # 不匹配，正常换行
            level = len(m.group(1))
            title = m.group(2).strip()
            if not title:
                return
            # 去掉 # 前缀，保留标题文字
            self.text.delete(f'{line_no}.0', f'{line_no}.end')
            self.text.insert(f'{line_no}.0', title)
            self.text.mark_set(tk.INSERT, f'{line_no}.end')
            # 记录目录条目
            offset = int(self.text.count('1.0', f'{line_no}.0', 'chars')[0])
            import uuid as _uuid
            entry = {
                'id':     str(_uuid.uuid4()),
                'level':  level,
                'text':   title[:60],
                'offset': offset,
            }
            entries = self._get_toc_entries(self._current_file)
            entries = [e for e in entries if e['offset'] != offset]
            entries.append(entry)
            self._save_toc_entries(self._current_file, entries)
            # 高亮标题
            self._apply_heading_tag(f'{line_no}.0', f'{line_no}.end', level)
            # 换行（手动插入，因为 return 'break' 会阻止默认行为）
            self.text.insert(tk.INSERT, '\n')
            # 刷新目录面板
            if self._sb_tab.get() == 'toc':
                self._refresh_toc()
            return 'break'
        except Exception:
            return   # 出错则正常换行

    def _autosave(self):
        self._autosave_after = None
        if not self._current_file:
            return
        ext = os.path.splitext(self._current_file)[1].lower()
        if ext in ('.pdf', '.docx', '.doc'):
            return
        try:
            content = self.text.get('1.0', tk.END)
            open(self._current_file, 'w', encoding='utf-8').write(content)
            # 顶栏短暂显示「已保存」提示
            fname = os.path.basename(self._current_file)
            self._lbl_file.config(text=f"{fname}  ✓")
            self.after(1500, lambda: self._lbl_file.config(text=fname))
        except Exception:
            pass

    # ── 搜索 ───────────────────────────────────────────────

    def _toggle_search(self):
        if self._search_bar.winfo_ismapped():
            self._close_search()
        else:
            self._search_bar.pack(fill=tk.X, before=self._vsb_text)
            self._search_divider.pack(fill=tk.X, before=self._vsb_text)
            self._search_entry.focus_set()
            self._search_entry.select_range(0, tk.END)

    def _close_search(self):
        self._search_bar.pack_forget()
        self._search_divider.pack_forget()
        self._clear_search_highlights()
        self._search_matches = []
        self._search_idx = -1
        self._search_count_lbl.config(text="")
        self.text.focus_set()

    def _search_changed(self, *_):
        if self._search_timer:
            self.after_cancel(self._search_timer)
        self._search_timer = self.after(300, self._do_search)

    def _do_search(self):
        self._clear_search_highlights()
        self._search_matches = []
        self._search_idx = -1
        kw = self._search_var.get()
        if not kw:
            self._search_count_lbl.config(text="")
            return
        idx = '1.0'
        while True:
            pos = self.text.search(kw, idx, nocase=True, stopindex=tk.END)
            if not pos:
                break
            end = f"{pos}+{len(kw)}c"
            self._search_matches.append((pos, end))
            self.text.tag_add('search_match', pos, end)
            idx = end
        total = len(self._search_matches)
        if total == 0:
            self._search_count_lbl.config(text="无结果")
            return
        self._search_idx = 0
        self._highlight_current()

    def _highlight_current(self):
        self.text.tag_remove('search_current', '1.0', tk.END)
        if not self._search_matches or self._search_idx < 0:
            self._search_count_lbl.config(text="")
            return
        pos, end = self._search_matches[self._search_idx]
        self.text.tag_add('search_current', pos, end)
        self.text.see(pos)
        total = len(self._search_matches)
        self._search_count_lbl.config(text=f"{self._search_idx + 1} / {total}")

    def _search_next(self):
        if not self._search_matches:
            return
        self._search_idx = (self._search_idx + 1) % len(self._search_matches)
        self._highlight_current()

    def _search_prev(self):
        if not self._search_matches:
            return
        self._search_idx = (self._search_idx - 1) % len(self._search_matches)
        self._highlight_current()

    def _clear_search_highlights(self):
        self.text.tag_remove('search_match', '1.0', tk.END)
        self.text.tag_remove('search_current', '1.0', tk.END)

    def _text_rclick(self, event):
        # 每次弹出前重建，确保标注选项是最新状态
        self._text_ctx.delete(0, tk.END)
        self._text_ctx.add_command(label="复制",
                                   command=lambda: self.text.event_generate('<<Copy>>'))
        self._text_ctx.add_command(label="粘贴",
                                   command=lambda: self.text.event_generate('<<Paste>>'))
        self._text_ctx.add_separator()
        self._text_ctx.add_command(label="撤销",
                                   command=lambda: self.text.edit_undo())
        self._text_ctx.add_command(label="重做",
                                   command=lambda: self.text.edit_redo())
        if self.annot_mgr:
            self.annot_mgr.populate_context_menu(self._text_ctx)

        # 目录标题标记（仅 txt 文件可手动标记）
        ext = os.path.splitext(self._current_file or '')[1].lower()
        if ext not in ('.pdf', '.docx', '.doc'):
            self._text_ctx.add_separator()
            self._text_ctx.add_command(label="取消目录标题", command=self._remove_heading)

        self._text_ctx.tk_popup(event.x_root, event.y_root)

    # ── 阅读计时 ───────────────────────────────────────────────

    @staticmethod
    def _today():
        from datetime import date
        return date.today().isoformat()  # "2026-04-08"

    @staticmethod
    def _fmt_secs(secs):
        if secs < 60:
            return f"{secs}s"
        elif secs < 3600:
            return f"{secs // 60}m {secs % 60:02d}s"
        else:
            h = secs // 3600
            m = (secs % 3600) // 60
            return f"{h}h {m:02d}m"

    def _on_focus_in(self, event):
        if event.widget is self and self._reading_start is None:
            import time
            self._reading_start = time.time()

    def _on_focus_out(self, event):
        if event.widget is self:
            self._flush_reading_time()

    def _flush_reading_time(self):
        import time
        if self._reading_start is None or not self._current_file:
            self._reading_start = None
            return
        elapsed = int(time.time() - self._reading_start)
        self._reading_start = None
        if elapsed < 2:
            return
        today = self._today()
        # 存储结构：{date: {filepath: seconds}}
        rt = self.store.files.get('__reading_time__', {})
        if not isinstance(rt, dict):
            rt = {}
        # 兼容旧版（{filepath: seconds}）→ 迁移到今日
        if rt and not any(isinstance(v, dict) for v in rt.values()):
            rt = {today: rt}
        day = rt.setdefault(today, {})
        day[self._current_file] = day.get(self._current_file, 0) + elapsed
        self.store.files['__reading_time__'] = rt
        self.store.save()
        self._update_time_label()

    def _update_time_label(self):
        """顶栏显示今日当前文件阅读时长"""
        if not self._current_file:
            self._lbl_time.config(text="")
            return
        today = self._today()
        rt    = self.store.files.get('__reading_time__', {})
        secs  = 0
        if isinstance(rt, dict):
            day_data = rt.get(today, {})
            if isinstance(day_data, dict):
                secs = day_data.get(self._current_file, 0)
        self._lbl_time.config(
            text=f"⏱ {self._fmt_secs(secs)}" if secs else "",
            cursor='hand2')
        self._lbl_time.bind('<Button-1>', lambda e: self._show_stats())

    def _show_stats(self):
        """弹出阅读统计窗口（今天/本周/本月/今年）"""
        from datetime import date, timedelta

        rt = self.store.files.get('__reading_time__', {})
        if not isinstance(rt, dict):
            rt = {}
        if rt and not any(isinstance(v, dict) for v in rt.values()):
            rt = {self._today(): rt}

        today = date.today()
        today_str = today.isoformat()

        # 本周（周一起）
        week_start = today - timedelta(days=today.weekday())
        week_dates = {(week_start + timedelta(days=i)).isoformat() for i in range(7)}

        # 本月
        month_prefix = today.strftime('%Y-%m')
        month_dates  = {d for d in rt if d.startswith(month_prefix)}

        # 今年
        year_prefix = str(today.year)
        year_dates  = {d for d in rt if d.startswith(year_prefix)}

        TABS = [
            ("今天", {today_str}),
            ("本周", week_dates),
            ("本月", month_dates),
            ("今年", year_dates),
        ]

        def aggregate(dates_set):
            """汇总指定日期集合内各文件时长"""
            result = {}
            for d, day_data in rt.items():
                if d in dates_set and isinstance(day_data, dict):
                    for fp, secs in day_data.items():
                        result[fp] = result.get(fp, 0) + secs
            return {fp: s for fp, s in result.items() if s > 0}

        # ── 窗口 ──
        win = tk.Toplevel(self)
        win.title("")
        win.configure(bg=C['bg_sidebar'])
        win.resizable(False, False)
        win.geometry("360x520")

        # ── 顶栏 ──
        hdr = tk.Frame(win, bg=C['bg_sidebar'], height=52)
        hdr.pack(fill=tk.X)
        hdr.pack_propagate(False)
        tk.Label(hdr, text="阅读统计", bg=C['bg_sidebar'], fg=C['fg'],
                 font=('PingFang SC', 13)).pack(side=tk.LEFT, padx=20)
        close_lbl = tk.Label(hdr, text="✕", bg=C['bg_sidebar'], fg='#333355',
                             font=('PingFang SC', 13), cursor='hand2')
        close_lbl.pack(side=tk.RIGHT, padx=16)
        close_lbl.bind('<Button-1>', lambda e: win.destroy())
        close_lbl.bind('<Enter>', lambda e: close_lbl.config(fg=C['fg_file']))
        close_lbl.bind('<Leave>', lambda e: close_lbl.config(fg='#333355'))
        tk.Frame(win, bg=C['border'], height=1).pack(fill=tk.X)

        # ── Tab 栏 ──
        tab_bar = tk.Frame(win, bg=C['bg_sidebar'])
        tab_bar.pack(fill=tk.X, padx=16, pady=(12, 0))
        tab_btns = []

        # ── 可滚动内容区 ──
        canvas = tk.Canvas(win, bg=C['bg_sidebar'], highlightthickness=0)
        canvas.pack(fill=tk.BOTH, expand=True, pady=(8, 0))
        inner = tk.Frame(canvas, bg=C['bg_sidebar'])
        inner_id = canvas.create_window((0, 0), window=inner, anchor='nw')
        inner.bind('<Configure>',
                   lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.bind('<Configure>',
                    lambda e: canvas.itemconfig(inner_id, width=e.width))

        def _scroll(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), 'units')
        canvas.bind('<MouseWheel>', _scroll)
        inner.bind('<MouseWheel>', _scroll)

        def show_tab(idx):
            # 更新 Tab 样式
            for i, btn in enumerate(tab_btns):
                if i == idx:
                    btn.config(fg=C['fg'], bg=C['accent_dim'],
                               relief='flat', bd=0)
                else:
                    btn.config(fg=C['fg_dim'], bg=C['bg_sidebar'],
                               relief='flat', bd=0)

            # 清空内容
            for w in inner.winfo_children():
                w.destroy()
            canvas.yview_moveto(0)

            _, dates_set = TABS[idx]
            data = aggregate(dates_set)

            if not data:
                tk.Label(inner, text="暂无记录\n\n开始阅读后会自动记录",
                         bg=C['bg_sidebar'], fg=C['fg_hint'],
                         font=('PingFang SC', 12),
                         justify=tk.CENTER).pack(pady=60)
                return

            total = sum(data.values())

            # 累计时长摘要
            summary = tk.Frame(inner, bg=C['bg_sidebar'])
            summary.pack(fill=tk.X, padx=16, pady=(16, 8))
            tk.Label(summary, text="累计阅读", bg=C['bg_sidebar'],
                     fg=C['fg_dim'], font=('PingFang SC', 11)).pack(side=tk.LEFT)
            tk.Label(summary, text=self._fmt_secs(total), bg=C['bg_sidebar'],
                     fg=C['accent'], font=('PingFang SC', 15, 'bold')).pack(side=tk.RIGHT)
            tk.Frame(inner, bg=C['border'], height=1).pack(
                fill=tk.X, padx=16, pady=(0, 8))

            # 各文件卡片
            max_secs = max(data.values())
            for fp, secs in sorted(data.items(), key=lambda x: x[1], reverse=True):
                card = tk.Frame(inner, bg=C['bg_input'], padx=12, pady=8)
                card.pack(fill=tk.X, padx=16, pady=2)

                name_row = tk.Frame(card, bg=C['bg_input'])
                name_row.pack(fill=tk.X)
                tk.Label(name_row,
                         text=self._trim_name(os.path.basename(fp), 24),
                         bg=C['bg_input'], fg=C['fg'],
                         font=('PingFang SC', 12)).pack(side=tk.LEFT)
                tk.Label(name_row, text=self._fmt_secs(secs),
                         bg=C['bg_input'], fg=C['accent'],
                         font=('PingFang SC', 12)).pack(side=tk.RIGHT)

                bar_bg = tk.Frame(card, bg=C['border'], height=3)
                bar_bg.pack(fill=tk.X, pady=(5, 0))
                frac = secs / max_secs
                tk.Frame(bar_bg, bg=C['accent'],
                         height=3).place(x=0, y=0, relwidth=frac)

        # 创建 Tab 按钮
        for i, (label, _) in enumerate(TABS):
            btn = tk.Label(tab_bar, text=label, bg=C['bg_sidebar'],
                           fg=C['fg_dim'], font=('PingFang SC', 12),
                           padx=14, pady=5, cursor='hand2')
            btn.pack(side=tk.LEFT, padx=(0, 2))
            btn.bind('<Button-1>', lambda e, idx=i: show_tab(idx))
            tab_btns.append(btn)

        show_tab(0)

    # ── 图标切换 ───────────────────────────────────────────

    def _load_icon(self):
        saved = self.store.files.get('__icon__', 'zz')
        self._switch_icon(saved, save=False)

    def _switch_icon(self, which, save=True):
        fname = 'logo2.png' if which == 'custom' else 'logo.png'
        path  = os.path.join(self._app_dir, fname)
        try:
            icon = tk.PhotoImage(file=path)
            self.iconphoto(True, icon)
            self._icon_ref = icon   # 防止 GC
        except Exception:
            pass
        if save:
            self.store.files['__icon__'] = which
            self.store.save()
            if hasattr(self, '_icon_var'):
                self._icon_var.set(which)

    # ── 主题切换 ───────────────────────────────────────────

    # ── 跟随系统主题 ───────────────────────────────────────

    @staticmethod
    def _get_system_theme():
        """检测 macOS 当前外观，返回 'dark' 或 'light'"""
        try:
            r = subprocess.run(
                ['defaults', 'read', '-g', 'AppleInterfaceStyle'],
                capture_output=True, text=True)
            return 'dark' if r.stdout.strip() == 'Dark' else 'light'
        except Exception:
            return 'dark'

    def _set_theme_manual(self, name):
        """手动切换主题，同时关闭跟随系统"""
        self._follow_sys.set(False)
        if self._sys_poll_id:
            self.after_cancel(self._sys_poll_id)
            self._sys_poll_id = None
        self._apply_theme(name)
        self.store.files['__follow_sys__'] = False
        self.store.save()

    def _toggle_follow_system(self):
        if self._follow_sys.get():
            # 开启：立即同步一次，然后开始轮询
            self.store.files['__follow_sys__'] = True
            self.store.save()
            self._sync_system_theme()
        else:
            # 关闭：停止轮询
            if self._sys_poll_id:
                self.after_cancel(self._sys_poll_id)
                self._sys_poll_id = None
            self.store.files['__follow_sys__'] = False
            self.store.save()

    def _sync_system_theme(self):
        """轮询系统主题，有变化时自动切换"""
        if not self._follow_sys.get():
            return
        sys_theme = self._get_system_theme()
        # 与当前实际显示的主题对比（而非存储值），避免启动时漏刷新
        if sys_theme != self._theme_var.get():
            self._apply_theme(sys_theme)
        # 每 2 秒检测一次
        self._sys_poll_id = self.after(2000, self._sync_system_theme)

    def _apply_theme(self, name):
        """切换护眼/白天模式：递归重色所有 tk 控件 + 刷新 ttk 样式 + 通知标注管理器"""
        old = dict(C)          # 切换前的旧色值快照
        _theme.apply(name)     # 更新全局 C（in-place）

        # 建立旧色 → 新色映射表
        color_map = {}
        for k in old:
            if old[k] != C[k]:
                color_map[old[k].lower()] = C[k]

        # 递归更新所有 tk 控件颜色
        self._recolor_widget(self, color_map)

        # 刷新 ttk Style（Treeview / Scrollbar）
        self._apply_ttk_styles()

        # 刷新搜索高亮 tag
        self.text.tag_configure('search_match',
                                background=C['hl_all_bg'], foreground=C['hl_all_fg'])
        self.text.tag_configure('search_current',
                                background=C['hl_cur_bg'], foreground=C['hl_cur_fg'])

        # 刷新侧边栏 Treeview 自定义 tag 颜色
        if hasattr(self, 'tree'):
            self.tree.tag_configure('tag',    foreground=C['fg_tag'])
            self.tree.tag_configure('file',   foreground=C['fg_file'])
            self.tree.tag_configure('header', foreground=C['fg_dim2'])

        # 刷新文本区颜色（不依赖 color_map，直接用新主题色值）
        self.text.configure(
            bg=C['bg'],
            fg=C['fg'],
            insertbackground=C['fg'],
            selectbackground=C['select_bg'],
            selectforeground=C['fg'],
        )

        # 通知标注管理器刷新
        if self.annot_mgr:
            self.annot_mgr.refresh_theme(color_map, theme_name=name)

        # 保存偏好
        self.store.files['__theme__'] = name
        self.store.save()

        # 更新菜单打勾状态
        self._theme_var.set(name)

    def _recolor_widget(self, widget, color_map):
        """递归将所有 tk 控件中与 color_map 匹配的颜色替换为新主题色"""
        PROPS = ('bg', 'fg', 'highlightbackground',
                 'insertbackground', 'selectbackground', 'activebackground',
                 'activeforeground', 'troughcolor')
        for w in widget.winfo_children():
            for prop in PROPS:
                try:
                    val = w.cget(prop)
                    # tkinter 返回的可能是 tuple，取最后一个元素
                    if isinstance(val, tuple):
                        val = val[-1]
                    if isinstance(val, str) and val.lower() in color_map:
                        w.configure(**{prop: color_map[val.lower()]})
                except Exception:
                    pass
            if w.winfo_children():
                self._recolor_widget(w, color_map)

    def _apply_ttk_styles(self):
        """重新应用 Treeview / Scrollbar 的 ttk.Style"""
        s = ttk.Style()
        s.configure('Sidebar.Treeview',
                    background=C['bg_sidebar'], foreground=C['fg_tag'],
                    fieldbackground=C['bg_sidebar'], rowheight=36,
                    borderwidth=0, indent=16)
        s.map('Sidebar.Treeview',
              background=[('selected', C['bg_sel_tag'])],
              foreground=[('selected', C['accent'])])
        s.configure('Thin.Vertical.TScrollbar',
                    troughcolor=C['bg'], background=C['accent_dim'],
                    width=4, relief=tk.FLAT, borderwidth=0)
        s.map('Thin.Vertical.TScrollbar',
              background=[('active', C['accent'])])

    # ── 目录导航 ───────────────────────────────────────────

    def _switch_sidebar_tab(self, tab):
        """切换侧栏：标签 / 目录"""
        self._sb_tab.set(tab)
        self._tab_lbl_tags._upd()
        self._tab_lbl_toc._upd()
        if tab == 'tags':
            self._tree_frame_ref.tkraise()
        else:
            self._toc_frame.configure(bg=C['bg_sidebar'])
            self._apply_ttk_styles()
            self._toc_frame.tkraise()
            self._refresh_toc()

    def _get_toc_entries(self, filepath):
        """从 store 取目录条目列表"""
        toc = self.store.files.get('__toc__', {})
        return toc.get(filepath, []) if isinstance(toc, dict) else []

    def _save_toc_entries(self, filepath, entries):
        toc = self.store.files.get('__toc__', {})
        if not isinstance(toc, dict):
            toc = {}
        toc[filepath] = entries
        self.store.files['__toc__'] = toc
        self.store.save()

    def _refresh_toc(self):
        """重建目录 Treeview"""
        # 保存当前展开状态
        self._save_toc_open_states()
        self._toc_tree.delete(*self._toc_tree.get_children())
        if not self._current_file:
            return
        open_states = self.store.files.get('__toc_open__', {})
        if not isinstance(open_states, dict):
            open_states = {}
        entries = self._get_toc_entries(self._current_file)
        if not entries:
            self._toc_tree.insert('', tk.END, text="  暂无目录条目",
                                  tags=('hint',))
            self._toc_tree.tag_configure('hint', foreground=C['fg_dim'])
            return
        stack = {}   # level -> iid
        for entry in sorted(entries, key=lambda x: x['offset']):
            lv = entry['level']
            parent = stack.get(lv - 1, '')
            indent = '  ' * (lv - 1)
            iid = self._toc_tree.insert(
                parent, tk.END,
                text=f"{indent}{entry['text']}",
                values=(entry['offset'], entry.get('id', '')),
                tags=(f'h{lv}',),
                open=open_states.get(entry.get('id', ''), False))
            stack[lv] = iid
            for deeper in range(lv + 1, 8):
                stack.pop(deeper, None)
        # 颜色配置
        for lv in range(1, 8):
            self._toc_tree.tag_configure(
                f'h{lv}', foreground=C['accent'] if lv == 1 else C['fg_tag'])
        # 点击后延迟 150ms 保存（覆盖展开/折叠两种动作，比虚拟事件更可靠）
        self._toc_tree.bind('<Button-1>', lambda e: self.after(150, self._save_toc_open_states))

    def _save_toc_open_states(self):
        """把目录 Treeview 当前的展开状态按 entry id 保存（树为空时跳过，避免覆盖历史）"""
        if not self._toc_tree.get_children(''):
            return
        states = self.store.files.get('__toc_open__', {})
        if not isinstance(states, dict):
            states = {}
        for iid in self._toc_tree.get_children(''):
            self._collect_toc_open(iid, states)
        self.store.files['__toc_open__'] = states
        self.store.save()

    def _collect_toc_open(self, iid, states):
        vals = self._toc_tree.item(iid, 'values')
        if vals and len(vals) >= 2:
            entry_id = vals[1]
            if entry_id:
                states[entry_id] = self._toc_tree.item(iid, 'open')
        for child in self._toc_tree.get_children(iid):
            self._collect_toc_open(child, states)

    def _on_toc_select(self, event):
        sel = self._toc_tree.focus()
        if not sel:
            return
        vals = self._toc_tree.item(sel, 'values')
        if not vals:
            return
        try:
            offset = int(vals[0])
            target = f'1.0+{offset}c'
            self.text.see(target)
            self.text.mark_set('insert', target)
            self.text.update_idletasks()
            self.text.yview(target)
        except Exception:
            pass

    def _mark_heading(self, level):
        """将选中文字标记为指定级别目录标题（txt）"""
        try:
            start_idx = self.text.index(tk.SEL_FIRST)
            end_idx   = self.text.index(tk.SEL_LAST)
            text_content = self.text.get(start_idx, end_idx).strip()
        except tk.TclError:
            # 无选中，取当前行
            insert = self.text.index(tk.INSERT)
            line_start = insert.split('.')[0] + '.0'
            line_end   = insert.split('.')[0] + '.end'
            start_idx  = line_start
            end_idx    = line_end
            text_content = self.text.get(start_idx, end_idx).strip()
        if not text_content or not self._current_file:
            return

        offset = int(self.text.count('1.0', start_idx, 'chars')[0])
        import uuid as _uuid
        entry = {
            'id':     str(_uuid.uuid4()),
            'level':  level,
            'text':   text_content[:60],
            'offset': offset,
        }
        entries = self._get_toc_entries(self._current_file)
        # 去重：同偏移量已有条目则替换
        entries = [e for e in entries if e['offset'] != offset]
        entries.append(entry)
        self._save_toc_entries(self._current_file, entries)

        # 高亮显示
        self._apply_heading_tag(start_idx, end_idx, level)

        if self._sb_tab.get() == 'toc':
            self._refresh_toc()
        else:
            self._switch_sidebar_tab('toc')

    def _remove_heading(self):
        """取消选中文字的标题标记"""
        try:
            start_idx = self.text.index(tk.SEL_FIRST)
        except tk.TclError:
            insert = self.text.index(tk.INSERT)
            start_idx = insert.split('.')[0] + '.0'
        if not self._current_file:
            return
        offset = int(self.text.count('1.0', start_idx, 'chars')[0])
        entries = self._get_toc_entries(self._current_file)
        # 移除最近的同行条目
        line = self.text.index(start_idx).split('.')[0]
        entries = [e for e in entries
                   if self.text.index(f'1.0+{e["offset"]}c').split('.')[0] != line]
        self._save_toc_entries(self._current_file, entries)
        self._refresh_toc()

    def _apply_heading_tag(self, start_idx, end_idx, level):
        """在 Text 控件中高亮显示标题标记"""
        tag = f'toc_h{level}'
        self.text.tag_add(tag, start_idx, end_idx)

    def _configure_heading_tags(self):
        """配置标题高亮 tag 样式"""
        colors = [C['accent'], C['fg'], C['fg'], C['fg_tag'],
                  C['fg_tag'], C['fg_dim'], C['fg_dim']]
        sizes  = [0, -1, -1, -2, -2, -3, -3]   # 相对于正文字号的偏移
        for lv in range(1, 8):
            tag = f'toc_h{lv}'
            sz  = max(10, self._font_size + sizes[lv - 1])
            weight = 'bold' if lv <= 2 else 'normal'
            self.text.tag_configure(tag,
                foreground=colors[lv - 1],
                font=(self._font_family, sz, weight))

    def _restore_heading_tags(self):
        """文件加载后恢复标题高亮"""
        if not self._current_file:
            return
        entries = self._get_toc_entries(self._current_file)
        for entry in entries:
            lv  = entry['level']
            off = entry['offset']
            end_off = off + len(entry['text'])
            try:
                si = f'1.0+{off}c'
                ei = f'1.0+{end_off}c'
                self._apply_heading_tag(si, ei, lv)
            except Exception:
                pass

    def _extract_word_toc(self, doc, filepath):
        """从 python-docx Document 提取标题，存入 store"""
        import re as _re
        entries = []
        char_offset = 0
        for para in doc.paragraphs:
            name = para.style.name
            sid  = getattr(para.style, 'style_id', '') or ''
            m = _re.match(r'Heading\s*(\d)', name) or _re.match(r'Heading(\d)', sid)
            if m and para.text.strip():
                lv = min(int(m.group(1)), 7)
                import uuid as _uuid
                entries.append({
                    'id':     str(_uuid.uuid4()),
                    'level':  lv,
                    'text':   para.text.strip()[:60],
                    'offset': char_offset,
                })
            char_offset += len(para.text) + 1   # +1 for \n
        if entries:
            self._save_toc_entries(filepath, entries)

    def _extract_pdf_toc(self, filepath):
        """从 PDF 提取书签目录（如有），存入 store"""
        try:
            import fitz
            doc = fitz.open(filepath)
            toc = doc.get_toc()   # [[level, title, page], ...]
            if not toc:
                return
            # 构建页码 → 字符偏移的映射
            page_offsets = []
            offset = 0
            for page in doc:
                page_offsets.append(offset)
                offset += len(page.get_text()) + 2
            import uuid as _uuid
            entries = []
            for level, title, page in toc:
                pg = max(0, page - 1)
                char_off = page_offsets[pg] if pg < len(page_offsets) else 0
                entries.append({
                    'id':     str(_uuid.uuid4()),
                    'level':  min(level, 7),
                    'text':   title.strip()[:60],
                    'offset': char_off,
                })
            if entries:
                self._save_toc_entries(filepath, entries)
        except Exception:
            pass

    def _open_global_search(self):
        """打开全局检索浮窗（已开则聚焦）"""
        if hasattr(self, '_global_search_win') and \
                self._global_search_win and \
                self._global_search_win.winfo_exists():
            self._global_search_win.lift()
            self._global_search_win.focus_force()
            return
        self._global_search_win = GlobalSearchWindow(self)

    def _on_file_drop(self, event):
        """处理拖放文件事件"""
        raw = event.data.strip()
        # tkinterdnd2 在 macOS 上多文件用空格分隔，路径含空格用 {} 包裹
        import re as _re
        paths = _re.findall(r'\{([^}]+)\}|(\S+)', raw)
        files = [p[0] or p[1] for p in paths]
        supported = ('.txt', '.pdf', '.docx', '.doc', '.zzg')
        valid = [f for f in files if os.path.isfile(f)
                 and os.path.splitext(f)[1].lower() in supported]
        if not valid:
            return
        # 导入所有文件到 store
        for fp in valid:
            self.store.add_file(fp)
        self.store.save()
        self._refresh_tree()
        # 加载第一个文件
        self._load(valid[0])

    def _show_help(self):
        win = tk.Toplevel(self)
        win.title("使用说明")
        win.resizable(False, False)
        win.configure(bg=C['bg'])
        win.transient(self)

        text = tk.Text(win, bg=C['bg'], fg=C['fg'],
                       font=(UI_FONT, 14), wrap=tk.WORD,
                       bd=0, highlightthickness=0,
                       padx=28, pady=24,
                       width=46, height=32,
                       state=tk.NORMAL, cursor='arrow')
        text.pack(fill=tk.BOTH, expand=True)

        HELP = """逐字稿 · 使用说明

━━ 导入文稿 ━━━━━━━━━━━━━━━━━

• 文件 → 导入 TXT / Word / PDF
• 或直接将文件拖入窗口

━━ 目录导航 ━━━━━━━━━━━━━━━━━

侧栏点击「目录」标签可查看目录。

Word / PDF 文件加载后自动提取目录。

TXT 文件手动标记标题：
  在行首输入 # 空格 标题文字，按 Enter

  例：# 第一章 引言       → 一级标题
      ## 1.1 背景         → 二级标题
      ### 小节            → 三级标题
      最多支持七级（####### ）

  规则：# 后必须有一个空格，
        空格后必须有文字，
        按 Enter 后前缀自动消失。

━━ 标注 ━━━━━━━━━━━━━━━━━━━━━

选中文字后：
• Cmd+1  黄色高亮
• Cmd+2  绿色高亮
• Cmd+3  红色高亮
• Cmd+4  下划线
• Cmd+5  批注

━━ 全局搜索 ━━━━━━━━━━━━━━━━━

Cmd+Shift+F 搜索所有文稿内容

━━ 其他快捷键 ━━━━━━━━━━━━━━━

Cmd+F       当前文稿内搜索
Cmd+\\      显示 / 隐藏侧栏
Cmd+=       放大字号
Cmd+−       缩小字号
"""
        text.insert('1.0', HELP)
        text.configure(state=tk.DISABLED)

        btn = tk.Button(win, text="关闭",
                        bg=C['bg_input'], fg=C['fg'],
                        activebackground=C['select_bg'],
                        activeforeground=C['fg'],
                        relief=tk.FLAT, bd=0,
                        font=(UI_FONT, 13),
                        padx=20, pady=8,
                        cursor='hand2',
                        command=win.destroy)
        btn.pack(pady=(0, 16))
        win.bind('<Escape>', lambda e: win.destroy())

        win.update_idletasks()
        x = self.winfo_x() + (self.winfo_width()  - win.winfo_width())  // 2
        y = self.winfo_y() + (self.winfo_height() - win.winfo_height()) // 2
        win.geometry(f'+{x}+{y}')
        win.lift()
        win.focus_force()

    def _on_close(self):
        self._flush_reading_time()
        self._save_position()
        self._save_toc_open_states()
        self.destroy()


def _read_file_safe(fp):
    """安全读取文本文件内容，失败返回空字符串"""
    ext = os.path.splitext(fp)[1].lower()
    if ext == '.zzg':
        try:
            import zipfile
            with zipfile.ZipFile(fp, 'r') as zf:
                return zf.read('content.txt').decode('utf-8')
        except Exception:
            return ''
    if ext in ('.docx', '.doc'):
        try:
            import docx
            doc = docx.Document(fp)
            return '\n'.join(p.text for p in doc.paragraphs)
        except Exception:
            return ''
    if ext == '.pdf':
        try:
            import fitz
            doc = fitz.open(fp)
            text = '\n\n'.join(page.get_text() for page in doc)
            doc.close()
            return text
        except Exception:
            return ''
    for enc in ('utf-8', 'gbk', 'utf-16'):
        try:
            return open(fp, encoding=enc).read()
        except Exception:
            continue
    return ''


class GlobalSearchWindow(tk.Toplevel):
    """全局跨文件检索浮窗——卡片式结果"""

    MAX_HITS   = 60
    CTX_BEFORE = 50
    CTX_AFTER  = 100
    FILE_LIMIT = 5

    def __init__(self, app: 'App'):
        super().__init__(app)
        self.app      = app
        self._timer   = None
        self._results = []
        self._cards   = []
        self._sel_idx = -1
        self._cache   = {}   # fp -> (mtime, content) 文件内容缓存
        self._configure_window()
        self._build()
        self.bind('<Escape>', lambda e: self._close())
        self.protocol('WM_DELETE_WINDOW', self._close)

    def _configure_window(self):
        self.title("全局搜索")
        self.configure(bg=C['bg'])
        self.geometry("760x620")
        self.minsize(560, 400)
        self.resizable(True, True)
        ax = self.app.winfo_x() + (self.app.winfo_width()  - 760) // 2
        ay = self.app.winfo_y() + (self.app.winfo_height() - 620) // 2
        self.geometry(f"760x620+{max(0,ax)}+{max(0,ay)}")

    def _build(self):
        # ── 搜索输入区 ──────────────────────────────────
        search_area = tk.Frame(self, bg=C['bg'], pady=20)
        search_area.pack(fill=tk.X, padx=32)

        wrap = tk.Frame(search_area, bg=C['bg_input'],
                        highlightbackground=C['accent'],
                        highlightthickness=2)
        wrap.pack(fill=tk.X)

        tk.Label(wrap, text="⌕", bg=C['bg_input'], fg=C['fg_dim'],
                 font=('PingFang SC', 18)).pack(side=tk.LEFT, padx=(14, 4))

        self._var = tk.StringVar()
        self._entry = tk.Entry(wrap, textvariable=self._var,
                               bg=C['bg_input'], fg=C['fg'],
                               insertbackground=C['fg'],
                               font=('PingFang SC', 17),
                               relief=tk.FLAT, bd=0)
        self._entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=12)
        self._entry.focus_set()

        self._count_lbl = tk.Label(wrap, text="", bg=C['bg_input'],
                                   fg=C['fg_dim'], font=('PingFang SC', 12))
        self._count_lbl.pack(side=tk.RIGHT, padx=14)

        self._var.trace_add('write', self._on_changed)
        self._entry.bind('<Return>', lambda e: self._open_sel())
        self._entry.bind('<Down>',   lambda e: (self._move_sel(1),  'break'))
        self._entry.bind('<Up>',     lambda e: (self._move_sel(-1), 'break'))

        # 分隔线
        tk.Frame(self, bg=C['border'], height=1).pack(fill=tk.X)

        # ── 结果滚动区 ──────────────────────────────────
        outer = tk.Frame(self, bg=C['bg'])
        outer.pack(fill=tk.BOTH, expand=True)

        vsb = tk.Scrollbar(outer, orient=tk.VERTICAL)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        self._canvas = tk.Canvas(outer, bg=C['bg'],
                                 highlightthickness=0,
                                 yscrollcommand=vsb.set)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=self._canvas.yview)

        self._list_frame = tk.Frame(self._canvas, bg=C['bg'])
        self._canvas_win = self._canvas.create_window(
            (0, 0), window=self._list_frame, anchor='nw')

        self._list_frame.bind('<Configure>', self._on_list_configure)
        self._canvas.bind('<Configure>',
            lambda e: self._canvas.itemconfig(self._canvas_win, width=e.width))

        # Toplevel 级别绑定：捕获 Frame/Label 上的滚轮（不消费事件的控件）
        self.bind('<MouseWheel>', self._on_canvas_scroll)

        self._show_hint("在已导入的所有文件中搜索")

    def _on_list_configure(self, _e):
        self._canvas.configure(
            scrollregion=(0, 0, 0, self._list_frame.winfo_reqheight()))

    def _on_canvas_scroll(self, e):
        self._canvas.yview_scroll(int(-1 * (e.delta / 120)), 'units')
        return 'break'   # 阻止 tk.Text 等控件消费事件

    # ── 搜索逻辑 ──────────────────────────────────────────

    def _on_changed(self, *_):
        if self._timer:
            self.after_cancel(self._timer)
        self._timer = self.after(150, self._do_search)

    def _get_cached(self, fp):
        """读取文件内容，结果按修改时间缓存，避免每次搜索重复 I/O"""
        try:
            mtime = os.path.getmtime(fp)
        except OSError:
            return ''
        if fp in self._cache and self._cache[fp][0] == mtime:
            return self._cache[fp][1]
        content = _read_file_safe(fp)
        self._cache[fp] = (mtime, content)
        return content

    def _do_search(self):
        query = self._var.get().strip()
        self._results.clear()
        self._sel_idx = -1

        if not query:
            self._show_hint("在已导入的所有文件中搜索")
            self._count_lbl.config(text="")
            return

        all_fps = [fp for fp in self.app.store.files
                   if not fp.startswith('__') and os.path.isfile(fp)]
        ql = query.lower()

        for fp in all_fps:
            content = self._get_cached(fp)
            if not content:
                continue
            cl = content.lower()
            pos, hits = 0, 0
            while hits < self.FILE_LIMIT and len(self._results) < self.MAX_HITS:
                idx = cl.find(ql, pos)
                if idx == -1:
                    break
                cs  = max(0, idx - self.CTX_BEFORE)
                ce  = min(len(content), idx + len(query) + self.CTX_AFTER)
                ctx = content[cs:ce].replace('\n', ' ')
                self._results.append((fp, idx, idx + len(query),
                                      ctx, idx - cs, idx - cs + len(query)))
                pos = idx + len(query)
                hits += 1

        self._render_results(query)

    # ── 渲染卡片 ──────────────────────────────────────────

    def _clear_list(self):
        for w in self._list_frame.winfo_children():
            w.destroy()
        self._cards.clear()

    def _show_hint(self, msg):
        self._clear_list()
        tk.Label(self._list_frame, text=msg,
                 bg=C['bg'], fg=C['fg_dim'],
                 font=('PingFang SC', 14)).pack(pady=60)

    def _render_results(self, query):
        self._clear_list()

        if not self._results:
            self._count_lbl.config(text="无结果")
            tk.Label(self._list_frame,
                     text=f'找不到 "{query}" 的相关内容',
                     bg=C['bg'], fg=C['fg_dim'],
                     font=('PingFang SC', 14)).pack(pady=60)
            return

        self._count_lbl.config(text=f"{len(self._results)} 条结果")

        for i, (fp, ms, me, ctx, hl_s, hl_e) in enumerate(self._results):
            self._make_card(i, fp, ms, ctx, hl_s, hl_e)

        self._canvas.yview_moveto(0)

    def _make_card(self, idx, fp, ms, ctx, hl_s, hl_e):
        BG     = C['bg_input']
        BG_HOV = C['hover_card']

        card = tk.Frame(self._list_frame, bg=BG,
                        cursor='hand2')
        card.pack(fill=tk.X, padx=24, pady=6)
        self._cards.append(card)

        inner = tk.Frame(card, bg=BG)
        inner.pack(fill=tk.X, padx=18, pady=14)

        # 文件名行
        fname = os.path.basename(fp)
        tk.Label(inner, text=f'📄  {fname}',
                 bg=BG, fg=C['accent'],
                 font=('PingFang SC', 13, 'bold'),
                 anchor='w').pack(fill=tk.X)

        # 分隔
        tk.Frame(inner, bg=C['border'], height=1).pack(fill=tk.X, pady=(6, 8))

        # 摘要：用小 Text 控件实现关键词高亮
        snippet = tk.Text(inner,
                          bg=BG, fg=C['fg'],
                          font=('PingFang SC', 15),
                          relief=tk.FLAT, bd=0,
                          highlightthickness=0,
                          wrap=tk.WORD,
                          cursor='hand2',
                          state=tk.NORMAL,
                          height=3)
        snippet.pack(fill=tk.X)

        snippet.tag_configure('hl',
            background=C['hl_all_bg'],
            foreground=C['hl_all_fg'])

        snippet.insert(tk.END, ctx[:hl_s])
        snippet.insert(tk.END, ctx[hl_s:hl_e], 'hl')
        snippet.insert(tk.END, ctx[hl_e:])
        snippet.config(state=tk.DISABLED)
        # tk.Text 会消费 MouseWheel，必须单独绑定转发给 canvas 并 return 'break'
        snippet.bind('<MouseWheel>', self._on_canvas_scroll)

        # 绑定点击和 hover（card、inner、snippet 三层都要绑）
        def _click(e, n=idx):
            self._open_result(n)

        def _enter(e, bg=BG_HOV, n=idx):
            self._set_card_bg(n, bg)

        def _leave(e, n=idx):
            if n == self._sel_idx:
                self._set_card_bg(n, C['select_bg'])
            else:
                self._set_card_bg(n, BG)

        for w in (card, inner, snippet):
            w.bind('<Button-1>', _click)
            w.bind('<Enter>',    _enter)
            w.bind('<Leave>',    _leave)
        for lbl in inner.winfo_children():
            if isinstance(lbl, (tk.Label, tk.Frame)):
                lbl.bind('<Button-1>', _click)
                lbl.bind('<Enter>',    _enter)
                lbl.bind('<Leave>',    _leave)

    def _set_card_bg(self, idx, color):
        if idx >= len(self._cards):
            return
        card = self._cards[idx]
        _repaint = [card]
        while _repaint:
            w = _repaint.pop()
            try:
                if not isinstance(w, tk.Scrollbar):
                    w.configure(bg=color)
            except Exception:
                pass
            _repaint.extend(w.winfo_children())

    # ── 键盘导航 ──────────────────────────────────────────

    def _move_sel(self, delta):
        if not self._cards:
            return
        if 0 <= self._sel_idx < len(self._cards):
            self._set_card_bg(self._sel_idx, C['bg_input'])
        self._sel_idx = (self._sel_idx + delta) % len(self._cards)
        self._set_card_bg(self._sel_idx, C['select_bg'])
        # 滚动到可见
        card = self._cards[self._sel_idx]
        card.update_idletasks()
        y = card.winfo_y()
        h = self._list_frame.winfo_reqheight()
        self._canvas.yview_moveto(y / h if h else 0)

    def _open_sel(self):
        if self._sel_idx >= 0:
            self._open_result(self._sel_idx)
        elif self._results:
            self._open_result(0)

    def _close(self):
        self.destroy()

    def _open_result(self, idx):
        if idx >= len(self._results):
            return
        fp, ms, me, *_ = self._results[idx]
        query = self._var.get().strip()
        self._close()
        self.app._load(fp)

        def jump():
            try:
                self.app.text.see(f'1.0+{ms}c')
                self.app.text.mark_set('insert', f'1.0+{ms}c')
            except Exception:
                pass
            try:
                if not self.app._search_bar.winfo_ismapped():
                    self.app._toggle_search()
                self.app._search_var.set(query)
                self.app._search_entry.select_range(0, tk.END)
            except Exception:
                pass

        self.app.after(350, jump)


if __name__ == '__main__':
    App().mainloop()
